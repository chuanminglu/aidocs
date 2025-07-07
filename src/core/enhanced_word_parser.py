"""
增强的Word文档解析器模块

新增功能：
1. 图片处理：支持Word文档中的图片提取和显示
2. 复杂表格：增强表格处理能力，支持合并单元格等
3. 样式保持：更好的格式和样式保持
"""

import os
import base64
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
import logging
import tempfile
import shutil
from io import BytesIO

# 图片处理
try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

# Word文档处理
try:
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph
    WORD_SUPPORT_AVAILABLE = True
except ImportError:
    WORD_SUPPORT_AVAILABLE = False

# XML解析
try:
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False

import zipfile

@dataclass
class ImageInfo:
    """图片信息"""
    filename: str
    width: Optional[int] = None
    height: Optional[int] = None
    format: Optional[str] = None
    base64_data: Optional[str] = None
    local_path: Optional[str] = None
    description: Optional[str] = None

@dataclass
class TableCellInfo:
    """表格单元格信息"""
    text: str
    row_span: int = 1
    col_span: int = 1
    alignment: str = "left"
    is_header: bool = False
    background_color: Optional[str] = None
    text_color: Optional[str] = None
    bold: bool = False
    italic: bool = False

@dataclass
class TableInfo:
    """表格信息"""
    rows: List[List[TableCellInfo]]
    alignment: str = "left"
    has_header: bool = False
    caption: Optional[str] = None

@dataclass
class StyleInfo:
    """样式信息"""
    font_name: Optional[str] = None
    font_size: Optional[int] = None
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: Optional[str] = None
    background_color: Optional[str] = None
    alignment: str = "left"
    indent: int = 0

@dataclass
class ParagraphInfo:
    """段落信息"""
    text: str
    style: StyleInfo = field(default_factory=StyleInfo)
    level: int = 0
    is_heading: bool = False
    bullet_level: Optional[int] = None
    numbering_level: Optional[int] = None

@dataclass
class EnhancedWordParseResult:
    """增强的Word解析结果"""
    success: bool
    content: str = ""
    markdown_content: str = ""
    paragraphs: List[ParagraphInfo] = field(default_factory=list)
    tables: List[TableInfo] = field(default_factory=list)
    images: List[ImageInfo] = field(default_factory=list)
    styles: Dict[str, StyleInfo] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    error_message: str = ""

class EnhancedWordParser:
    """增强的Word文档解析器"""
    
    def __init__(self, extract_images: bool = True, preserve_styles: bool = True):
        self.logger = logging.getLogger(__name__)
        self.extract_images = extract_images
        self.preserve_styles = preserve_styles
        self.temp_dir = None
        self._check_dependencies()
    
    def _check_dependencies(self) -> bool:
        """检查依赖是否可用"""
        if not WORD_SUPPORT_AVAILABLE:
            self.logger.error("Word处理依赖不可用")
            return False
        if self.extract_images and not PILLOW_AVAILABLE:
            self.logger.warning("图片处理依赖不可用，图片功能将被禁用")
            self.extract_images = False
        return True
    
    def __enter__(self):
        """上下文管理器入口"""
        if self.extract_images:
            self.temp_dir = tempfile.mkdtemp(prefix="aidocs_word_")
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """上下文管理器退出"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def parse_document(self, file_path: str) -> EnhancedWordParseResult:
        """解析Word文档"""
        if not WORD_SUPPORT_AVAILABLE:
            return EnhancedWordParseResult(
                success=False,
                error_message="Word支持库未安装"
            )
        
        if not os.path.exists(file_path):
            return EnhancedWordParseResult(
                success=False,
                error_message=f"文件不存在: {file_path}"
            )
        
        try:
            # 打开文档
            doc = Document(file_path)
            
            # 提取基本信息
            paragraphs = self._extract_paragraphs(doc)
            tables = self._extract_tables(doc)
            images = self._extract_images(file_path, doc) if self.extract_images else []
            styles = self._extract_styles(doc) if self.preserve_styles else {}
            metadata = self._extract_metadata(doc)
            
            # 生成内容
            content = self._generate_plain_text(paragraphs, tables)
            markdown_content = self._generate_markdown(paragraphs, tables, images)
            
            return EnhancedWordParseResult(
                success=True,
                content=content,
                markdown_content=markdown_content,
                paragraphs=paragraphs,
                tables=tables,
                images=images,
                styles=styles,
                metadata=metadata
            )
            
        except Exception as e:
            self.logger.error(f"解析Word文档失败: {e}")
            return EnhancedWordParseResult(
                success=False,
                error_message=f"解析失败: {str(e)}"
            )
    
    def _extract_paragraphs(self, doc: DocxDocument) -> List[ParagraphInfo]:
        """提取段落信息"""
        paragraphs = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # 获取段落样式
            style = self._get_paragraph_style(para)
            
            # 判断是否为标题
            is_heading = False
            level = 0
            if para.style.name.startswith('Heading'):
                is_heading = True
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
            
            # 检查编号和项目符号
            bullet_level = None
            numbering_level = None
            if para._element.pPr is not None:
                numPr = para._element.pPr.numPr
                if numPr is not None:
                    if numPr.ilvl is not None:
                        bullet_level = int(numPr.ilvl.val)
                    if numPr.numId is not None:
                        numbering_level = int(numPr.numId.val)
            
            paragraph_info = ParagraphInfo(
                text=para.text,
                style=style,
                level=level,
                is_heading=is_heading,
                bullet_level=bullet_level,
                numbering_level=numbering_level
            )
            paragraphs.append(paragraph_info)
        
        return paragraphs
    
    def _extract_tables(self, doc: DocxDocument) -> List[TableInfo]:
        """提取表格信息"""
        tables = []
        
        for table in doc.tables:
            table_info = self._parse_table(table)
            tables.append(table_info)
        
        return tables
    
    def _parse_table(self, table) -> TableInfo:
        """解析单个表格"""
        rows = []
        has_header = False
        
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            
            for cell_idx, cell in enumerate(row.cells):
                # 获取单元格信息
                cell_info = self._parse_table_cell(cell)
                
                # 检查是否为表头
                if row_idx == 0 and self._is_header_cell(cell):
                    cell_info.is_header = True
                    has_header = True
                
                row_cells.append(cell_info)
            
            rows.append(row_cells)
        
        # 获取表格对齐方式
        alignment = self._get_table_alignment(table)
        
        return TableInfo(
            rows=rows,
            alignment=alignment,
            has_header=has_header
        )
    
    def _parse_table_cell(self, cell: _Cell) -> TableCellInfo:
        """解析表格单元格"""
        text = cell.text.strip()
        
        # 获取合并信息
        row_span = 1
        col_span = 1
        
        # 检查垂直合并
        if cell._element.tcPr is not None:
            vMerge = cell._element.tcPr.vMerge
            if vMerge is not None:
                if vMerge.val == 'restart':
                    row_span = self._count_merged_rows(cell)
                elif vMerge.val is None:  # 继续合并
                    row_span = 0  # 标记为被合并的单元格
        
        # 检查水平合并
        if cell._element.tcPr is not None:
            gridSpan = cell._element.tcPr.gridSpan
            if gridSpan is not None:
                col_span = int(gridSpan.val)
        
        # 获取单元格样式
        alignment = self._get_cell_alignment(cell)
        background_color = self._get_cell_background_color(cell)
        text_color = self._get_cell_text_color(cell)
        bold = self._is_cell_bold(cell)
        italic = self._is_cell_italic(cell)
        
        return TableCellInfo(
            text=text,
            row_span=row_span,
            col_span=col_span,
            alignment=alignment,
            background_color=background_color,
            text_color=text_color,
            bold=bold,
            italic=italic
        )
    
    def _extract_images(self, file_path: str, doc: DocxDocument) -> List[ImageInfo]:
        """提取图片信息"""
        if not self.extract_images or not PILLOW_AVAILABLE:
            return []
        
        images = []
        
        try:
            # 使用zipfile提取图片
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # 获取所有图片文件
                image_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                
                for img_file in image_files:
                    try:
                        # 提取图片数据
                        img_data = zip_file.read(img_file)
                        
                        # 保存到临时文件
                        img_filename = os.path.basename(img_file)
                        if self.temp_dir:
                            local_path = os.path.join(self.temp_dir, img_filename)
                            with open(local_path, 'wb') as f:
                                f.write(img_data)
                        else:
                            local_path = None
                        
                        # 获取图片信息
                        with Image.open(BytesIO(img_data)) as img:
                            width, height = img.size
                            format = img.format
                        
                        # 转换为base64
                        base64_data = base64.b64encode(img_data).decode('utf-8')
                        
                        image_info = ImageInfo(
                            filename=img_filename,
                            width=width,
                            height=height,
                            format=format,
                            base64_data=base64_data,
                            local_path=local_path
                        )
                        images.append(image_info)
                        
                    except Exception as e:
                        self.logger.warning(f"处理图片 {img_file} 失败: {e}")
                        continue
        
        except Exception as e:
            self.logger.error(f"提取图片失败: {e}")
        
        return images
    
    def _extract_styles(self, doc: DocxDocument) -> Dict[str, StyleInfo]:
        """提取样式信息"""
        styles = {}
        
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_info = self._parse_style(style)
                styles[style.name] = style_info
        
        return styles
    
    def _extract_metadata(self, doc: DocxDocument) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision,
                'version': core_props.version or ''
            })
        except Exception as e:
            self.logger.warning(f"提取元数据失败: {e}")
        
        return metadata
    
    def _get_paragraph_style(self, para: Paragraph) -> StyleInfo:
        """获取段落样式"""
        style = StyleInfo()
        
        # 获取段落级别样式
        if para.style.font.name:
            style.font_name = para.style.font.name
        if para.style.font.size:
            style.font_size = int(para.style.font.size.pt)
        if para.style.font.bold:
            style.bold = para.style.font.bold
        if para.style.font.italic:
            style.italic = para.style.font.italic
        if para.style.font.underline:
            style.underline = True
        
        # 获取段落对齐方式
        if para.alignment:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
            }
            style.alignment = alignment_map.get(para.alignment, "left")
        
        # 获取缩进
        if para.paragraph_format.left_indent:
            style.indent = int(para.paragraph_format.left_indent.pt)
        
        return style
    
    def _get_table_alignment(self, table) -> str:
        """获取表格对齐方式"""
        try:
            if table.alignment == WD_TABLE_ALIGNMENT.LEFT:
                return "left"
            elif table.alignment == WD_TABLE_ALIGNMENT.CENTER:
                return "center"
            elif table.alignment == WD_TABLE_ALIGNMENT.RIGHT:
                return "right"
        except Exception:
            pass
        return "left"
    
    def _is_header_cell(self, cell: _Cell) -> bool:
        """判断是否为表头单元格"""
        # 检查是否有特殊样式或格式
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.bold:
                        return True
        except Exception:
            pass
        return False
    
    def _count_merged_rows(self, cell: _Cell) -> int:
        """计算合并的行数"""
        # 这是一个简化的实现
        # 实际的合并行数计算可能更复杂
        return 1
    
    def _get_cell_alignment(self, cell: _Cell) -> str:
        """获取单元格对齐方式"""
        try:
            for paragraph in cell.paragraphs:
                if paragraph.alignment:
                    alignment_map = {
                        WD_ALIGN_PARAGRAPH.LEFT: "left",
                        WD_ALIGN_PARAGRAPH.CENTER: "center",
                        WD_ALIGN_PARAGRAPH.RIGHT: "right",
                        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
                    }
                    return alignment_map.get(paragraph.alignment, "left")
        except Exception:
            pass
        return "left"
    
    def _get_cell_background_color(self, cell: _Cell) -> Optional[str]:
        """获取单元格背景色"""
        try:
            shading = cell._element.tcPr.shd
            if shading is not None and shading.fill:
                return f"#{shading.fill}"
        except Exception:
            pass
        return None
    
    def _get_cell_text_color(self, cell: _Cell) -> Optional[str]:
        """获取单元格文本颜色"""
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.color and run.font.color.rgb:
                        return f"#{run.font.color.rgb}"
        except Exception:
            pass
        return None
    
    def _is_cell_bold(self, cell: _Cell) -> bool:
        """判断单元格是否为粗体"""
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.bold:
                        return True
        except Exception:
            pass
        return False
    
    def _is_cell_italic(self, cell: _Cell) -> bool:
        """判断单元格是否为斜体"""
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.italic:
                        return True
        except Exception:
            pass
        return False
    
    def _parse_style(self, style) -> StyleInfo:
        """解析样式"""
        style_info = StyleInfo()
        
        try:
            if style.font.name:
                style_info.font_name = style.font.name
            if style.font.size:
                style_info.font_size = int(style.font.size.pt)
            if style.font.bold:
                style_info.bold = style.font.bold
            if style.font.italic:
                style_info.italic = style.font.italic
            if style.font.underline:
                style_info.underline = True
        except Exception:
            pass
        
        return style_info
    
    def _generate_plain_text(self, paragraphs: List[ParagraphInfo], tables: List[TableInfo]) -> str:
        """生成纯文本内容"""
        content = []
        
        for para in paragraphs:
            content.append(para.text)
        
        for table in tables:
            content.append("")  # 空行
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row)
                content.append(row_text)
        
        return "\n".join(content)
    
    def _generate_markdown(self, paragraphs: List[ParagraphInfo], tables: List[TableInfo], images: List[ImageInfo]) -> str:
        """生成Markdown内容"""
        markdown_lines = []
        
        # 处理段落
        for para in paragraphs:
            if para.is_heading:
                # 标题
                markdown_lines.append(f"{'#' * para.level} {para.text}")
            elif para.bullet_level is not None:
                # 项目符号
                indent = "  " * para.bullet_level
                markdown_lines.append(f"{indent}- {para.text}")
            elif para.numbering_level is not None:
                # 编号列表
                indent = "  " * (para.numbering_level - 1) if para.numbering_level > 1 else ""
                markdown_lines.append(f"{indent}1. {para.text}")
            else:
                # 普通段落
                if para.text.strip():
                    markdown_lines.append(para.text)
            
            markdown_lines.append("")  # 空行
        
        # 处理表格
        for table in tables:
            markdown_lines.append("")  # 空行
            
            if table.rows:
                # 表格头部
                if table.has_header:
                    header_row = table.rows[0]
                    header_text = " | ".join(cell.text for cell in header_row)
                    markdown_lines.append(f"| {header_text} |")
                    
                    # 分隔线
                    separator = " | ".join("---" for _ in header_row)
                    markdown_lines.append(f"| {separator} |")
                    
                    # 数据行
                    for row in table.rows[1:]:
                        row_text = " | ".join(cell.text for cell in row)
                        markdown_lines.append(f"| {row_text} |")
                else:
                    # 没有表头的表格
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row)
                        markdown_lines.append(f"| {row_text} |")
            
            markdown_lines.append("")  # 空行
        
        # 处理图片
        for image in images:
            markdown_lines.append("")  # 空行
            alt_text = image.description or f"图片: {image.filename}"
            
            if image.local_path:
                # 使用本地路径
                markdown_lines.append(f"![{alt_text}]({image.local_path})")
            elif image.base64_data:
                # 使用base64数据
                data_url = f"data:image/{image.format.lower()};base64,{image.base64_data}"
                markdown_lines.append(f"![{alt_text}]({data_url})")
            
            markdown_lines.append("")  # 空行
        
        return "\n".join(markdown_lines)
