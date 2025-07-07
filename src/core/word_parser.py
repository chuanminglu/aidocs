"""
Word文档解析器模块

提供Word文档（.docx/.doc）的读取、解析、转换和保存功能。
此模块完全独立于主编辑器，确保Word功能的隔离性和可维护性。

主要功能：
1. Word文档读取和文本提取
2. Word文档结构解析（标题层级、段落、表格等）
3. Word转Markdown转换
4. Markdown转Word保存
5. Word文档大纲结构提取
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from dataclasses import dataclass
import logging

# Word文档处理依赖
try:
    from docx import Document
    from docx.document import Document as DocxDocument
    import docx2txt
    WORD_SUPPORT_AVAILABLE = True
except ImportError:
    WORD_SUPPORT_AVAILABLE = False
    logging.warning("Word文档支持库未安装，Word功能将不可用")
    # 创建占位符类型，避免类型错误
    DocxDocument = Any

# 其他依赖
import zipfile
import xml.etree.ElementTree as ET
import tempfile
import shutil

@dataclass
class OutlineItem:
    """大纲项目数据结构"""
    text: str
    level: int
    line_number: int
    item_type: str = "heading"

@dataclass
class WordParseResult:
    """Word解析结果"""
    success: bool
    content: str = ""
    outline: Optional[List[OutlineItem]] = None
    metadata: Optional[Dict[str, Any]] = None
    error_message: str = ""
    
    def __post_init__(self):
        if self.outline is None:
            self.outline = []
        if self.metadata is None:
            self.metadata = {}
        if self.metadata is None:
            self.metadata = {}

class WordDocumentParser:
    """Word文档解析器
    
    功能隔离设计：
    - 此类独立处理所有Word相关操作
    - 不依赖主编辑器和其他GUI组件
    - 提供统一的接口供外部调用
    - 错误处理不影响主应用程序
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self._check_dependencies()
    
    def _check_dependencies(self) -> bool:
        """检查Word处理依赖是否可用"""
        if not WORD_SUPPORT_AVAILABLE:
            self.logger.error("Word处理依赖不可用，请安装 python-docx 和 docx2txt")
            return False
        return True
    
    def is_word_file(self, file_path: str) -> bool:
        """检查文件是否为Word文档"""
        if not file_path:
            return False
        
        suffix = Path(file_path).suffix.lower()
        return suffix in ['.docx', '.doc']
    
    def extract_text_content(self, file_path: str) -> WordParseResult:
        """提取Word文档的纯文本内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            WordParseResult: 解析结果，包含文本内容和元数据
        """
        if not WORD_SUPPORT_AVAILABLE:
            return WordParseResult(
                success=False,
                error_message="Word支持库未安装"
            )
        
        if not os.path.exists(file_path):
            return WordParseResult(
                success=False,
                error_message=f"文件不存在: {file_path}"
            )
        
        try:
            # 使用docx2txt快速提取纯文本
            text_content = docx2txt.process(file_path)
            
            # 获取文档元数据
            doc = Document(file_path)
            metadata = self._extract_metadata(doc)
            
            return WordParseResult(
                success=True,
                content=text_content,
                metadata=metadata
            )
            
        except Exception as e:
            self.logger.error(f"提取Word文档文本失败: {e}")
            return WordParseResult(
                success=False,
                error_message=f"解析失败: {str(e)}"
            )
    
    def extract_structured_content(self, file_path: str) -> WordParseResult:
        """提取Word文档的结构化内容
        
        解析Word文档的段落、标题、表格等结构，
        并转换为Markdown格式
        
        Args:
            file_path: Word文档路径
            
        Returns:
            WordParseResult: 包含Markdown内容和大纲结构
        """
        if not WORD_SUPPORT_AVAILABLE:
            return WordParseResult(
                success=False,
                error_message="Word支持库未安装"
            )
        
        # 首先尝试结构化解析
        try:
            doc = Document(file_path)
            markdown_content = []
            outline_items = []
            line_number = 1
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if not para_text:
                    markdown_content.append("")
                    line_number += 1
                    continue
                
                # 检查是否为标题
                try:
                    style_name = paragraph.style.name
                    if style_name and style_name.startswith('Heading'):
                        level = self._extract_heading_level(style_name)
                        if level > 0:
                            # 添加到大纲
                            outline_items.append(OutlineItem(
                                text=para_text,
                                level=level,
                                line_number=line_number,
                                item_type="heading"
                            ))
                            # 转换为Markdown标题
                            markdown_content.append('#' * level + ' ' + para_text)
                        else:
                            markdown_content.append(para_text)
                    else:
                        # 普通段落
                        markdown_content.append(para_text)
                except Exception:
                    # 如果样式访问失败，作为普通段落处理
                    markdown_content.append(para_text)
                
                line_number += 1
            
            # 处理表格
            try:
                for table in doc.tables:
                    table_md = self._convert_table_to_markdown(table)
                    markdown_content.extend(table_md)
                    line_number += len(table_md)
            except Exception as e:
                self.logger.warning(f"表格处理失败，跳过: {e}")
            
            # 获取元数据
            try:
                metadata = self._extract_metadata(doc)
            except Exception:
                metadata = {"title": Path(file_path).stem}
            
            return WordParseResult(
                success=True,
                content='\n\n'.join(markdown_content),
                outline=outline_items,
                metadata=metadata
            )
            
        except Exception as e:
            self.logger.error(f"结构化解析失败: {e}")
            # 如果结构化解析失败，尝试简单文本提取
            return self._fallback_text_extraction(file_path, str(e))
    
    def get_outline_items(self, file_path: str) -> List[OutlineItem]:
        """从Word文档提取大纲结构
        
        专门用于大纲导航功能
        
        Args:
            file_path: Word文档路径
            
        Returns:
            List[OutlineItem]: 大纲项目列表
        """
        result = self.extract_structured_content(file_path)
        if result.success and result.outline:
            return result.outline
        return []
    
    def convert_to_markdown(self, file_path: str) -> str:
        """将Word文档转换为Markdown格式
        
        Args:
            file_path: Word文档路径
            
        Returns:
            str: Markdown内容，失败时返回空字符串
        """
        result = self.extract_structured_content(file_path)
        if result.success:
            return result.content
        else:
            self.logger.error(f"Word转Markdown失败: {result.error_message}")
            return ""
    
    def save_as_word(self, content: str, output_path: str, 
                     title: Optional[str] = None) -> bool:
        """将Markdown内容保存为Word文档
        
        Args:
            content: Markdown内容
            output_path: 输出文件路径
            title: 文档标题
            
        Returns:
            bool: 保存成功返回True
        """
        if not WORD_SUPPORT_AVAILABLE:
            self.logger.error("Word支持库未安装，无法保存Word文档")
            return False
        
        try:
            doc = Document()
            
            # 设置文档标题
            if title:
                doc.core_properties.title = title
            
            # 解析Markdown并转换为Word格式
            self._convert_markdown_to_word(doc, content)
            
            # 保存文档
            doc.save(output_path)
            self.logger.info(f"Word文档已保存到: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"保存Word文档失败: {e}")
            return False
    
    def _convert_markdown_to_word(self, doc: Any, content: str) -> None:
        """将Markdown内容转换为Word文档格式"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            self._process_markdown_line(doc, line)
    
    def _process_markdown_line(self, doc: Any, line: str) -> None:
        """处理单行Markdown内容"""
        # 处理标题
        if line.startswith('#'):
            self._add_heading_to_doc(doc, line)
        # 处理列表
        elif line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            doc.add_paragraph(list_text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(list_text, style='List Number')
        # 普通段落
        else:
            doc.add_paragraph(line)
    
    def _add_heading_to_doc(self, doc: Any, line: str) -> None:
        """添加标题到Word文档"""
        level = len(line) - len(line.lstrip('#'))
        title_text = line.lstrip('#').strip()
        
        if level <= 6:  # Word最多支持6级标题
            doc.add_heading(title_text, level)
        else:
            # 超过6级的标题作为普通段落处理
            doc.add_paragraph(title_text)
    
    def _extract_heading_level(self, style_name: str) -> int:
        """从样式名称提取标题级别"""
        if style_name.startswith('Heading'):
            try:
                # 尝试从样式名中提取数字
                level_str = re.search(r'\d+', style_name)
                if level_str:
                    return int(level_str.group())
            except ValueError:
                pass
        return 0
    
    def _extract_metadata(self, doc: Any) -> Dict[str, Any]:
        """提取Word文档元数据"""
        try:
            core_props = doc.core_properties
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'category': core_props.category or '',
                'comments': core_props.comments or '',
                'keywords': core_props.keywords or '',
                'language': core_props.language or '',
                'revision': core_props.revision or 0
            }
        except Exception as e:
            self.logger.warning(f"提取元数据失败: {e}")
            return {}
    
    def _convert_table_to_markdown(self, table) -> List[str]:
        """将Word表格转换为Markdown表格"""
        try:
            markdown_table = []
            
            # 处理表格行
            for i, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                markdown_row = '| ' + ' | '.join(row_cells) + ' |'
                markdown_table.append(markdown_row)
                
                # 添加表头分隔符
                if i == 0:
                    separator = '| ' + ' | '.join(['---'] * len(row_cells)) + ' |'
                    markdown_table.append(separator)
            
            # 添加空行分隔
            markdown_table.append('')
            return markdown_table
            
        except Exception as e:
            self.logger.warning(f"转换表格失败: {e}")
            return [f"[表格转换失败: {str(e)}]", '']
    
    def _fallback_text_extraction(self, file_path: str, original_error: str) -> WordParseResult:
        """备用文本提取方法
        
        当结构化解析失败时，使用多种方法进行文本提取
        
        Args:
            file_path: Word文档路径
            original_error: 原始错误信息
            
        Returns:
            WordParseResult: 简单文本提取结果
        """
        self.logger.warning(f"使用备用方法提取文本: {original_error}")
        
        # 尝试多种备用方法
        fallback_methods = [
            self._try_docx2txt_extraction,
            self._try_zipfile_extraction,
            self._try_binary_extraction
        ]
        
        for method in fallback_methods:
            try:
                result = method(file_path, original_error)
                if result.success:
                    return result
            except Exception as e:
                self.logger.warning(f"备用方法 {method.__name__} 失败: {e}")
                continue
        
        # 所有方法都失败
        return WordParseResult(
            success=False,
            error_message=f"所有解析方法均失败: 原始错误={original_error}"
        )
    
    def _try_docx2txt_extraction(self, file_path: str, original_error: str) -> WordParseResult:
        """尝试使用docx2txt提取"""
        try:
            text_content = docx2txt.process(file_path)
            
            if not text_content or not text_content.strip():
                raise Exception("提取的内容为空")
            
            return self._process_extracted_text(text_content, file_path, "docx2txt", original_error)
            
        except Exception as e:
            raise Exception(f"docx2txt提取失败: {e}")
    
    def _try_zipfile_extraction(self, file_path: str, original_error: str) -> WordParseResult:
        """尝试直接解压Word文档XML"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # 尝试读取document.xml
                try:
                    doc_xml = zip_file.read('word/document.xml')
                    root = ET.fromstring(doc_xml)
                    
                    # 提取文本内容
                    text_elements = root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    text_content = '\n'.join([elem.text or '' for elem in text_elements if elem.text])
                    
                    if text_content.strip():
                        return self._process_extracted_text(text_content, file_path, "zipfile-xml", original_error)
                        
                except Exception:
                    # 如果document.xml不存在或损坏，尝试其他文件
                    for file_name in zip_file.namelist():
                        if file_name.endswith('.xml') and 'word' in file_name:
                            try:
                                xml_content = zip_file.read(file_name)
                                # 简单提取XML中的文本
                                import re
                                text = re.sub(r'<[^>]+>', ' ', xml_content.decode('utf-8', errors='ignore'))
                                text = re.sub(r'\s+', ' ', text).strip()
                                
                                if len(text) > 100:  # 如果提取到足够的文本
                                    return self._process_extracted_text(text, file_path, "zipfile-raw", original_error)
                            except Exception:
                                continue
            
            raise Exception("无法从ZIP结构中提取文本")
            
        except Exception as e:
            raise Exception(f"ZIP文件解析失败: {e}")
    
    def _try_binary_extraction(self, file_path: str, original_error: str) -> WordParseResult:
        """尝试二进制文本提取（最后的备用方案）"""
        try:
            import re
            
            with open(file_path, 'rb') as f:
                binary_data = f.read()
            
            # 尝试多种编码方式提取可读文本
            extracted_texts = []
            
            # 方法1: UTF-8解码
            try:
                text = binary_data.decode('utf-8', errors='ignore')
                if text.strip():
                    extracted_texts.append(('utf-8', text))
            except Exception:
                pass
            
            # 方法2: GBK解码（适合中文）
            try:
                text = binary_data.decode('gbk', errors='ignore')
                if text.strip():
                    extracted_texts.append(('gbk', text))
            except Exception:
                pass
            
            # 方法3: CP936解码（Windows中文）
            try:
                text = binary_data.decode('cp936', errors='ignore')
                if text.strip():
                    extracted_texts.append(('cp936', text))
            except Exception:
                pass
            
            # 方法4: Latin-1解码（保证不出错）
            try:
                text = binary_data.decode('latin-1', errors='ignore')
                if text.strip():
                    extracted_texts.append(('latin-1', text))
            except Exception:
                pass
            
            # 选择最好的解码结果
            best_text = ""
            best_encoding = "unknown"
            max_chinese_chars = 0
            
            for encoding, text in extracted_texts:
                # 统计中文字符数量
                chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
                if chinese_chars > max_chinese_chars:
                    max_chinese_chars = chinese_chars
                    best_text = text
                    best_encoding = encoding
            
            if not best_text:
                # 如果没有找到好的文本，使用第一个
                if extracted_texts:
                    best_encoding, best_text = extracted_texts[0]
                else:
                    raise Exception("无法从任何编码中提取文本")
            
            # 清理和提取有意义的文本
            cleaned_text = self._clean_binary_text(best_text)
            
            if len(cleaned_text.strip()) < 50:
                raise Exception("提取的文本内容太少")
            
            return self._process_extracted_text(cleaned_text, file_path, f"binary-{best_encoding}", original_error)
            
        except Exception as e:
            raise Exception(f"二进制提取失败: {e}")
    
    def _clean_binary_text(self, text: str) -> str:
        """清理从二进制数据中提取的文本"""
        import re
        
        # 移除控制字符和特殊符号，但保留中文、英文、数字和基本标点
        cleaned = re.sub(r'[^\u4e00-\u9fff\w\s\.\,\!\?\;\:\-\(\)\[\]\"\'""''（）《》【】、。，！？；：\n\r]+', ' ', text)
        
        # 合并多个空白字符
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        # 分割成行，过滤掉太短的行
        lines = []
        for line in cleaned.split('\n'):
            line = line.strip()
            if len(line) > 3:  # 只保留长度大于3的行
                lines.append(line)
        
        # 进一步处理：移除重复的行和明显的垃圾数据
        unique_lines = []
        seen = set()
        
        for line in lines:
            # 跳过重复行
            if line in seen:
                continue
            
            # 跳过全是重复字符的行
            if len(set(line.replace(' ', ''))) <= 2:
                continue
                
            # 跳过明显的编码错误行（连续特殊字符）
            if re.search(r'[^\u4e00-\u9fff\w\s]{5,}', line):
                continue
            
            # 检查是否包含有意义的内容
            meaningful_chars = len(re.findall(r'[\u4e00-\u9fff\w]', line))
            if meaningful_chars >= 3:  # 至少包含3个有意义字符
                unique_lines.append(line)
                seen.add(line)
        
        # 限制行数，避免太长
        if len(unique_lines) > 200:
            unique_lines = unique_lines[:200]
        
        return '\n'.join(unique_lines)
    
    def _process_extracted_text(self, text_content: str, file_path: str, method: str, original_error: str) -> WordParseResult:
        """处理提取的文本内容"""
        try:
            # 清理和分段
            lines = text_content.split('\n')
            processed_lines = []
            outline_items = []
            
            for i, line in enumerate(lines, 1):
                line = line.strip()
                if not line or len(line) < 2:
                    continue
                
                # 简单启发式标题检测
                if self._is_likely_heading(line):
                    level = self._guess_heading_level(line)
                    outline_items.append(OutlineItem(
                        text=line,
                        level=level,
                        line_number=len(processed_lines) + 1,
                        item_type="heading"
                    ))
                    processed_lines.append('#' * level + ' ' + line)
                else:
                    processed_lines.append(line)
            
            # 如果没有识别到标题，尝试基于长度和位置推断
            if not outline_items and processed_lines:
                self._infer_structure(processed_lines, outline_items)
            
            final_content = '\n\n'.join(processed_lines) if processed_lines else text_content
            
            return WordParseResult(
                success=True,
                content=final_content,
                outline=outline_items,
                metadata={
                    "title": Path(file_path).stem, 
                    "method": f"fallback-{method}",
                    "original_error": original_error,
                    "extracted_lines": len(processed_lines)
                },
                error_message=f"使用备用方法提取 ({method}，原因: {original_error[:100]}...)"
            )
            
        except Exception as e:
            raise Exception(f"文本处理失败: {e}")
    
    def _infer_structure(self, lines: list, outline_items: list):
        """基于文本特征推断文档结构"""
        try:
            for i, line in enumerate(lines):
                # 基于长度推断：较短的行可能是标题
                if len(line) < 50 and len(line) > 5:
                    # 检查是否包含关键词
                    if any(keyword in line for keyword in [
                        '课程', '培训', '第', '章', '节', '部分', '天', '日', '阶段',
                        '内容', '目标', '技术', '实战', '案例', '项目', '方案'
                    ]):
                        level = 2 if len(line) < 20 else 3
                        outline_items.append(OutlineItem(
                            text=line,
                            level=level,
                            line_number=i + 1,
                            item_type="heading"
                        ))
                        # 更新原line为标题格式
                        lines[i] = '#' * level + ' ' + line
                        
        except Exception:
            pass  # 推断失败不影响主要功能
    
    def _is_likely_heading(self, line: str) -> bool:
        """判断文本行是否可能是标题"""
        if not line or len(line) > 100:  # 标题通常不会太长
            return False
        
        # 启发式规则
        # 1. 全部大写（但排除全是英文字母的情况）
        if line.isupper() and len(line) > 2 and not line.isalpha():
            return True
        
        # 2. 包含数字编号
        if re.match(r'^[0-9]+[\.\)、]\s*', line):
            return True
            
        # 3. 包含中文编号
        if re.match(r'^[一二三四五六七八九十]+[\.\)、]\s*', line):
            return True
        
        # 4. 较短且以常见标题词开头
        heading_keywords = [
            '第', '章', '节', '部分', '概述', '总结', '介绍', '背景', 
            '方法', '结果', '结论', '目标', '任务', '计划', '方案',
            '课程', '培训', '学习', '实战', '案例', '项目', '技术'
        ]
        if len(line) < 50 and any(keyword in line for keyword in heading_keywords):
            return True
        
        # 5. 被特殊字符包围
        if re.match(r'^[\*\#\=\-]{2,}.*[\*\#\=\-]{2,}$', line):
            return True
            
        return False
    
    def _guess_heading_level(self, line: str) -> int:
        """猜测标题级别"""
        # 根据数字编号猜测
        if re.match(r'^[0-9]+[\.\)、]\s*', line):
            return 1
        
        # 根据中文编号猜测
        if re.match(r'^[一二三四五六七八九十]+[\.\)、]\s*', line):
            return 2
        
        # 根据关键词猜测
        if any(keyword in line for keyword in ['第', '章']):
            return 1
        elif any(keyword in line for keyword in ['节', '部分', '课程']):
            return 2
        elif any(keyword in line for keyword in ['目标', '任务', '方案']):
            return 3
        else:
            return 3
# Word功能可用性检查函数
def check_word_support() -> Tuple[bool, str]:
    """检查Word功能支持状态
    
    Returns:
        Tuple[bool, str]: (是否支持, 状态消息)
    """
    if WORD_SUPPORT_AVAILABLE:
        return True, "Word功能已启用"
    else:
        return False, "Word功能未启用，请安装依赖: pip install python-docx docx2txt"

# 便捷函数
def quick_word_to_markdown(file_path: str) -> str:
    """快速将Word文档转换为Markdown
    
    便捷函数，供外部快速调用
    """
    parser = WordDocumentParser()
    return parser.convert_to_markdown(file_path)

def quick_markdown_to_word(content: str, output_path: str, title: Optional[str] = None) -> bool:
    """快速将Markdown保存为Word文档
    
    便捷函数，供外部快速调用
    """
    parser = WordDocumentParser()
    return parser.save_as_word(content, output_path, title)
