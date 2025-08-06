#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
格式转换器模块 - MD2DOC项目
===========================

该模块负责将Markdown格式转换为Word文档格式，支持完整的格式映射。

主要功能：
- Markdown到Word格式映射
- 嵌套格式处理（粗体+斜体等）
- 链接、图片、代码格式转换
- 图表渲染结果集成
- 表格和列表格式转换

Author: MD2DOC Team
Date: 2025-08-04
Version: 1.0.0
"""

import re
import logging
import tempfile
import hashlib
from typing import Dict, List, Any, Optional, Tuple, TYPE_CHECKING
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

if TYPE_CHECKING:
    from docx import Document as DocxDocument
    from docx.document import Document as DocxDocumentType
else:
    DocxDocument = Any
    DocxDocumentType = Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError as e:
    logging.error(f"python-docx库未安装: {e}")
    raise ImportError("请安装python-docx库: pip install python-docx")


class FormatType(Enum):
    """格式类型枚举"""
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    BOLD = "bold"
    ITALIC = "italic"
    CODE_INLINE = "code_inline"
    CODE_BLOCK = "code_block"
    LINK = "link"
    IMAGE = "image"
    LIST_ORDERED = "list_ordered"
    LIST_UNORDERED = "list_unordered"
    TABLE = "table"
    BLOCKQUOTE = "blockquote"
    HORIZONTAL_RULE = "horizontal_rule"
    CHART = "chart"


@dataclass
class FormatRule:
    """格式转换规则"""
    markdown_pattern: str
    word_style: str
    processor_func: str
    priority: int = 0
    nested_support: bool = False


@dataclass
class TextFragment:
    """文本片段数据类"""
    content: str
    format_type: FormatType
    attributes: Dict[str, Any]
    start_pos: int = 0
    end_pos: int = 0


class FormatConverter:
    """格式转换器核心类"""
    
    def __init__(self, document: Optional[Any] = None):
        """初始化格式转换器
        
        Args:
            document: Word文档对象，如果为None则创建新文档
        """
        self.document = document or Document()
        self.logger = logging.getLogger(__name__)
        
        # 立即设置文档字体（必须在其他操作之前）
        self._force_set_document_font()
        
        # 初始化格式映射表
        self._init_format_rules()
        self._init_word_styles()
        
        # 格式处理缓存
        self._format_cache = {}
        self._nested_format_stack = []
        
    def _force_set_document_font(self):
        """强制设置文档字体为微软雅黑（包括目录）"""
        try:
            # 方法1：设置Normal样式的字体
            normal_style = self.document.styles['Normal']
            normal_style.font.name = 'Microsoft YaHei'
            normal_style.font.size = Pt(12)
            
            # 设置所有标题样式的字体
            for level in range(1, 10):  # 扩展到更多级别
                try:
                    heading_style = self.document.styles[f'Heading {level}']
                    heading_style.font.name = 'Microsoft YaHei'
                except KeyError:
                    pass
            
            # 设置目录相关样式的字体
            toc_styles = [
                'TOC Heading', 'TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'TOC 5', 'TOC 6',
                'TOC 7', 'TOC 8', 'TOC 9', 'Table of Contents', 'Contents Heading',
                'Contents 1', 'Contents 2', 'Contents 3', 'Contents 4', 'Contents 5'
            ]
            
            for style_name in toc_styles:
                try:
                    toc_style = self.document.styles[style_name]
                    toc_style.font.name = 'Microsoft YaHei'
                except KeyError:
                    pass  # 样式不存在，忽略
            
            # 方法2：尝试XML级别设置（更强制）
            try:
                from docx.oxml.ns import qn
                
                # 获取Normal样式的XML元素
                style_element = normal_style.element
                
                # 设置字符属性
                rPr = style_element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                
                if rFonts is None:
                    from docx.oxml import parse_xml
                    rFonts = parse_xml(f'<w:rFonts {rPr.nsmap} />')
                    rPr.append(rFonts)
                
                # 设置所有字体属性
                rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
                rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
                rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                rFonts.set(qn('w:cs'), 'Microsoft YaHei')
                
                self.logger.info("成功设置文档XML级别字体为微软雅黑")
                
            except Exception as xml_error:
                self.logger.warning(f"XML级别字体设置失败，但基础设置已完成: {xml_error}")
            
            # 方法3：强制设置文档级别的默认字体
            try:
                # 获取文档的settings部分
                document_xml = self.document._element
                settings = document_xml.find(qn('w:settings'))
                if settings is not None:
                    # 添加默认字体设置
                    from docx.oxml import parse_xml
                    default_fonts = parse_xml(f'''
                        <w:defaultFonts {document_xml.nsmap}
                            w:ascii="Microsoft YaHei"
                            w:eastAsia="Microsoft YaHei"
                            w:hAnsi="Microsoft YaHei"
                            w:cs="Microsoft YaHei"/>
                    ''')
                    settings.append(default_fonts)
                    
                self.logger.info("成功设置文档级别默认字体")
                
            except Exception as doc_error:
                self.logger.warning(f"文档级别字体设置失败: {doc_error}")
            
            self.logger.info("文档字体设置完成（包括目录样式）")
            
        except Exception as e:
            self.logger.error(f"字体设置失败: {e}")
            # 最基础的备用方案
            try:
                self.document.styles['Normal'].font.name = 'Microsoft YaHei'
            except Exception:
                pass
        
    def _init_format_rules(self):
        """初始化格式转换规则表"""
        self.format_rules = {
            FormatType.HEADING: FormatRule(
                markdown_pattern=r'^(#{1,6})\s+(.+)$',
                word_style='heading',
                processor_func='_process_heading',
                priority=10
            ),
            FormatType.BOLD: FormatRule(
                markdown_pattern=r'\*\*(.+?)\*\*',
                word_style='bold',
                processor_func='_process_bold',
                priority=5,
                nested_support=True
            ),
            FormatType.ITALIC: FormatRule(
                markdown_pattern=r'\*(.+?)\*',
                word_style='italic',
                processor_func='_process_italic',
                priority=4,
                nested_support=True
            ),
            FormatType.CODE_INLINE: FormatRule(
                markdown_pattern=r'`([^`]+)`',
                word_style='code_inline',
                processor_func='_process_code_inline',
                priority=6
            ),
            FormatType.CODE_BLOCK: FormatRule(
                markdown_pattern=r'```(\w+)?\n(.*?)\n```',
                word_style='code_block',
                processor_func='_process_code_block',
                priority=8
            ),
            FormatType.LINK: FormatRule(
                markdown_pattern=r'\[([^\]]+)\]\(([^)]+)\)',
                word_style='hyperlink',
                processor_func='_process_link',
                priority=7
            ),
            FormatType.IMAGE: FormatRule(
                markdown_pattern=r'!\[([^\]]*)\]\(([^)]+)\)',
                word_style='image',
                processor_func='_process_image',
                priority=9
            ),
            FormatType.LIST_UNORDERED: FormatRule(
                markdown_pattern=r'^[\s]*[-*+]\s+(.+)$',
                word_style='list_bullet',
                processor_func='_process_unordered_list',
                priority=6
            ),
            FormatType.LIST_ORDERED: FormatRule(
                markdown_pattern=r'^[\s]*\d+\.\s+(.+)$',
                word_style='list_number',
                processor_func='_process_ordered_list',
                priority=6
            ),
            FormatType.BLOCKQUOTE: FormatRule(
                markdown_pattern=r'^>\s+(.+)$',
                word_style='quote',
                processor_func='_process_blockquote',
                priority=5
            ),
            FormatType.TABLE: FormatRule(
                markdown_pattern=r'^\|(.+)\|$',
                word_style='table',
                processor_func='_process_table',
                priority=8
            ),
            FormatType.HORIZONTAL_RULE: FormatRule(
                markdown_pattern=r'^-{3,}$',
                word_style='horizontal_rule',
                processor_func='_process_horizontal_rule',
                priority=3
            )
        }
        
    def _init_word_styles(self):
        """初始化Word文档样式"""
        styles = self.document.styles
        
        # 设置默认字体为微软雅黑
        self._set_default_font()
        
        # 创建代码样式
        try:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
            code_font = code_style.font
            code_font.name = 'Consolas'  # 代码保持Consolas字体
            code_font.size = Pt(10)
            code_font.color.rgb = RGBColor(0x1f, 0x1f, 0x1f)
        except ValueError:
            pass  # 样式可能已存在
            
        # 创建代码块样式
        try:
            code_block_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_block_style.base_style = styles['Normal']
            code_block_font = code_block_style.font
            code_block_font.name = 'Consolas'  # 代码保持Consolas字体
            code_block_font.size = Pt(9)
            code_block_style.paragraph_format.left_indent = Inches(0.5)
            code_block_style.paragraph_format.space_before = Pt(6)
            code_block_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            pass
            
        # 创建引用样式
        try:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.base_style = styles['Normal']
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.space_before = Pt(6)
            quote_style.paragraph_format.space_after = Pt(6)
            quote_font = quote_style.font
            quote_font.name = 'Microsoft YaHei'  # 引用文本使用微软雅黑
            quote_font.italic = True
            quote_font.color.rgb = RGBColor(0x5a, 0x5a, 0x5a)
        except ValueError:
            pass
            
        # 创建表格样式
        self._init_table_styles()
    
    def _set_default_font(self):
        """设置文档默认字体为微软雅黑"""
        try:
            # 设置Normal样式的字体
            normal_style = self.document.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Microsoft YaHei'
            normal_font.size = Pt(12)
            
            # 设置标题样式的字体
            for level in range(1, 7):
                try:
                    heading_style = self.document.styles[f'Heading {level}']
                    heading_font = heading_style.font
                    heading_font.name = 'Microsoft YaHei'
                except KeyError:
                    pass
                    
        except Exception as e:
            self.logger.warning(f"设置默认字体失败: {e}")
    
    def _clean_text(self, text: str) -> str:
        """清理文本中的特殊字符和问题字符（保守版本 + Markdown语法清理）
        
        Args:
            text: 需要清理的文本
            
        Returns:
            清理后的文本
        """
        if not text:
            return text
            
        # 移除常见的问题字符
        cleaned_text = text
        
        # 清理Markdown语法标记（但保留内容，保留下划线）
        # 只处理星号标记的粗体和斜体
        cleaned_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned_text)  # **粗体**
        cleaned_text = re.sub(r'\*([^*\n]+)\*', r'\1', cleaned_text)    # *斜体*
        # 不处理下划线标记，保留原始下划线字符
        
        # 处理内联代码标记
        cleaned_text = re.sub(r'`([^`]+)`', r'\1', cleaned_text)        # `代码`
        
        # 处理链接标记但保留文本
        cleaned_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned_text)  # [文本](链接)
        
        # 只清理真正有害的特殊字符，不清理正常内容
        
        # 移除零宽字符和变体选择器
        cleaned_text = re.sub(r'[\u200B-\u200F]', '', cleaned_text)  # 零宽字符和格式字符
        cleaned_text = re.sub(r'[\uFEFF]', '', cleaned_text)  # 字节顺序标记
        cleaned_text = re.sub(r'[\uFE0E\uFE0F]', '', cleaned_text)  # 变体选择器
        cleaned_text = re.sub(r'\u180E', '', cleaned_text)  # 蒙古语空格
        cleaned_text = re.sub(r'[\u2060-\u206F]', '', cleaned_text)  # 其他特殊空格和格式字符
        
        # 移除控制字符（但保留换行符和制表符）
        cleaned_text = re.sub(r'[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F]', '', cleaned_text)
        
        # 移除双向格式字符
        cleaned_text = re.sub(r'[\u202A-\u202E]', '', cleaned_text)  # 双向格式字符
        cleaned_text = re.sub(r'[\u2066-\u2069]', '', cleaned_text)  # 双向隔离字符
        
        # 移除一些明确的问题字符
        cleaned_text = re.sub(r'[\u00AD]', '', cleaned_text)  # 软连字符
        cleaned_text = re.sub(r'[\u034F]', '', cleaned_text)  # 组合石墨烯连接符
        
        # 标准化空白字符（轻微处理）
        cleaned_text = re.sub(r'[\s]+', ' ', cleaned_text)  # 多个空格变为单个空格
        cleaned_text = cleaned_text.strip()  # 移除首尾空格
        
        return cleaned_text

    def _apply_font_to_run(self, run, font_name: str = 'Microsoft YaHei'):
        """为run对象应用指定字体（强制设置）
        
        Args:
            run: Word文档中的run对象
            font_name: 字体名称，默认为微软雅黑
        """
        try:
            # 强制设置字体的所有属性
            run.font.name = font_name
            
            # 设置XML级别的字体属性，确保真正生效
            from docx.oxml import parse_xml
            from docx.oxml.ns import qn
            
            # 获取run的XML元素
            r_pr = run._element.get_or_add_rPr()
            
            # 设置字体元素
            r_fonts = r_pr.find(qn('w:rFonts'))
            if r_fonts is None:
                r_fonts = parse_xml(f'<w:rFonts {r_pr.nsmap} />')
                r_pr.append(r_fonts)
            
            # 设置所有字体属性
            r_fonts.set(qn('w:ascii'), font_name)
            r_fonts.set(qn('w:hAnsi'), font_name)
            r_fonts.set(qn('w:eastAsia'), font_name)
            r_fonts.set(qn('w:cs'), font_name)
            
        except Exception as e:
            self.logger.warning(f"强制字体设置失败: {e}")
            # 简单备用方案
            try:
                run.font.name = font_name
            except Exception:
                pass
    
    def _create_run_with_font(self, paragraph, text: str = "") -> object:
        """创建新的run并自动应用微软雅黑字体
        
        Args:
            paragraph: 段落对象
            text: 运行文本内容
            
        Returns:
            新创建的run对象
        """
        run = paragraph.add_run(text)
        self._apply_font_to_run(run, 'Microsoft YaHei')
        return run
    
    def _init_table_styles(self):
        """初始化表格相关样式"""
        try:
            # 创建表格标题样式
            table_title_style = self.document.styles.add_style('TableTitle', WD_STYLE_TYPE.PARAGRAPH)
            table_title_style.base_style = self.document.styles['Normal']
            title_font = table_title_style.font
            title_font.name = 'Microsoft YaHei'  # 使用微软雅黑
            title_font.bold = True
            title_font.size = Pt(11)
            table_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table_title_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            pass
    
    def convert_markdown_to_word(self, markdown_text: str) -> Any:
        """将Markdown文本转换为Word文档
        
        Args:
            markdown_text: Markdown格式的文本
            
        Returns:
            转换后的Word文档对象
        """
        self.logger.info("开始Markdown到Word格式转换")
        
        # 预处理：分离图表代码块
        processed_text, chart_blocks = self._extract_chart_blocks(markdown_text)
        
        # 按行处理
        lines = processed_text.split('\n')
        current_table = []
        in_code_block = False
        code_block_content = []
        code_language = ""
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 处理代码块
            if line.strip().startswith('```'):
                if not in_code_block:
                    # 开始代码块
                    in_code_block = True
                    code_language = line.strip()[3:]
                    code_block_content = []
                else:
                    # 结束代码块
                    in_code_block = False
                    self._process_code_block('\n'.join(code_block_content), code_language)
                    code_block_content = []
                    code_language = ""
                i += 1
                continue
                
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
            
            # 处理表格
            if self._is_table_line(line):
                current_table.append(line)
                i += 1
                continue
            elif current_table:
                # 表格结束
                self._process_table_block(current_table)
                current_table = []
            
            # 处理普通行
            self._process_line(line, chart_blocks)
            i += 1
        
        # 处理剩余的表格
        if current_table:
            self._process_table_block(current_table)
            
        self.logger.info("Markdown到Word格式转换完成")
        return self.document
    
    def _extract_chart_blocks(self, text: str) -> Tuple[str, Dict[str, Dict[str, str]]]:
        """提取图表代码块
        
        Args:
            text: 原始Markdown文本
            
        Returns:
            (处理后的文本, 图表代码块字典)
        """
        chart_blocks: Dict[str, Dict[str, str]] = {}
        chart_pattern = r'```(mermaid|plantuml|chart)\n(.*?)\n```'
        
        def replace_chart(match):
            chart_type = match.group(1)
            chart_code = match.group(2)
            chart_id = f"chart_{len(chart_blocks)}"
            chart_blocks[chart_id] = {
                'type': chart_type,
                'code': chart_code
            }
            return f"{{{{CHART:{chart_id}}}}}"
        
        processed_text = re.sub(chart_pattern, replace_chart, text, flags=re.DOTALL)
        return processed_text, chart_blocks
    
    def _is_table_line(self, line: str) -> bool:
        """判断是否为表格行"""
        stripped = line.strip()
        return stripped.startswith('|') and stripped.endswith('|')
    
    def _process_line(self, line: str, chart_blocks: Dict[str, Dict[str, str]]):
        """处理单行文本"""
        line = line.strip()
        if not line:
            self._add_paragraph("")
            return
            
        # 检查是否包含图表占位符
        chart_match = re.search(r'\{\{CHART:(.+?)\}\}', line)
        if chart_match:
            chart_id = chart_match.group(1)
            if chart_id in chart_blocks:
                self._process_chart(chart_blocks[chart_id])
                return
        
        # 检查标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            self._process_heading(title, level)
            return
        
        # 检查列表
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            content = list_match.group(3)
            if marker.isdigit() or marker.endswith('.'):
                self._process_ordered_list(content, indent)
            else:
                self._process_unordered_list(content, indent)
            return
        
        # 检查引用
        quote_match = re.match(r'^>\s+(.+)$', line)
        if quote_match:
            self._process_blockquote(quote_match.group(1))
            return
        
        # 检查水平分割线
        if re.match(r'^-{3,}$', line):
            self._process_horizontal_rule()
            return
        
        # 处理普通段落（包含内联格式）
        self._process_paragraph_with_inline_formats(line)
    
    def _process_heading(self, title: str, level: int):
        """处理标题（强制应用微软雅黑字体）"""
        cleaned_title = self._clean_text(title)  # 清理标题文本
        heading = self.document.add_heading(cleaned_title, level)
        
        # 强制设置标题的字体为微软雅黑
        for paragraph in heading.runs if hasattr(heading, 'runs') else []:
            self._apply_font_to_run(paragraph, 'Microsoft YaHei')
        
        # 如果heading是paragraph对象，处理其runs
        if hasattr(heading, 'runs'):
            for run in heading.runs:
                self._apply_font_to_run(run, 'Microsoft YaHei')
        
        self.logger.debug(f"添加标题: 级别{level}, 内容: {cleaned_title}, 已应用微软雅黑字体")
    
    def _process_paragraph_with_inline_formats(self, text: str):
        """处理包含内联格式的段落"""
        paragraph = self.document.add_paragraph()
        
        # 处理内联格式的优先级顺序
        patterns = [
            (r'!\[([^\]]*)\]\(([^)]+)\)', self._add_image_to_paragraph),
            (r'\[([^\]]+)\]\(([^)]+)\)', self._add_link_to_paragraph),
            (r'`([^`]+)`', self._add_code_to_paragraph),
            (r'\*\*(.+?)\*\*', self._add_bold_to_paragraph),
            (r'\*(.+?)\*', self._add_italic_to_paragraph),
        ]
        
        current_text = text
        processed_ranges = []
        
        for pattern, processor in patterns:
            matches = list(re.finditer(pattern, current_text))
            for match in matches:
                start, end = match.span()
                # 检查是否与已处理的范围重叠
                if not any(s <= start < e or s < end <= e for s, e in processed_ranges):
                    processor(paragraph, match)
                    processed_ranges.append((start, end))
        
        # 添加剩余的普通文本
        self._add_remaining_text(paragraph, current_text, processed_ranges)
    
    def _add_image_to_paragraph(self, paragraph, match):
        """在段落中添加图片"""
        alt_text = match.group(1)
        image_path = match.group(2)
        
        try:
            # 检查图片路径是否存在
            if Path(image_path).exists():
                run = paragraph.add_run()
                # 为普通图片也使用智能尺寸计算
                optimal_width = self._calculate_optimal_image_width(image_path, "image")
                run.add_picture(image_path, width=optimal_width)
                self.logger.debug(f"添加图片: {image_path}, 宽度: {optimal_width.inches:.1f}英寸")
            else:
                # 添加图片占位符
                cleaned_alt_text = self._clean_text(alt_text)  # 清理图片alt文本
                run = paragraph.add_run(f"[图片: {cleaned_alt_text}]")
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                self._apply_font_to_run(run)  # 应用微软雅黑字体
                self.logger.warning(f"图片文件不存在: {image_path}")
        except Exception as e:
            self.logger.error(f"添加图片失败: {e}")
            cleaned_alt_text = self._clean_text(alt_text)  # 清理图片alt文本
            run = paragraph.add_run(f"[图片加载失败: {cleaned_alt_text}]")
            run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
            self._apply_font_to_run(run)  # 应用微软雅黑字体
    
    def _add_link_to_paragraph(self, paragraph, match):
        """在段落中添加链接"""
        link_text = self._clean_text(match.group(1))  # 清理文本
        url = match.group(2)
        
        # 添加超链接
        self._create_hyperlink(paragraph, url, link_text)
        self.logger.debug(f"添加链接: {link_text} -> {url}")
    
    def _add_code_to_paragraph(self, paragraph, match):
        """在段落中添加内联代码"""
        code_text = self._clean_text(match.group(1))  # 清理文本
        run = paragraph.add_run(code_text)
        try:
            run.style = self.document.styles['Code']
        except KeyError:
            # 如果样式不存在，使用基本格式
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        self.logger.debug(f"添加内联代码: {code_text}")
    
    def _add_bold_to_paragraph(self, paragraph, match):
        """在段落中添加粗体文本"""
        bold_text = self._clean_text(match.group(1))  # 清理文本
        run = paragraph.add_run(bold_text)
        run.bold = True
        self._apply_font_to_run(run)  # 应用微软雅黑字体
        self.logger.debug(f"添加粗体: {bold_text}")
    
    def _add_italic_to_paragraph(self, paragraph, match):
        """在段落中添加斜体文本"""
        italic_text = self._clean_text(match.group(1))  # 清理文本
        run = paragraph.add_run(italic_text)
        run.italic = True
        self._apply_font_to_run(run)  # 应用微软雅黑字体
        self.logger.debug(f"添加斜体: {italic_text}")
    
    def _add_remaining_text(self, paragraph, original_text: str, processed_ranges: List[Tuple[int, int]]):
        """添加剩余的普通文本"""
        # 按位置排序处理过的范围
        processed_ranges.sort()
        
        last_end = 0
        for start, end in processed_ranges:
            # 添加范围前的文本
            if start > last_end:
                text_before = self._clean_text(original_text[last_end:start])  # 清理文本
                if text_before.strip():
                    run = paragraph.add_run(text_before)
                    self._apply_font_to_run(run)  # 应用微软雅黑字体
            last_end = end
        
        # 添加最后的文本
        if last_end < len(original_text):
            remaining_text = self._clean_text(original_text[last_end:])  # 清理文本
            if remaining_text.strip():
                run = paragraph.add_run(remaining_text)
                self._apply_font_to_run(run)  # 应用微软雅黑字体
    
    def _create_hyperlink(self, paragraph, url: str, text: str):
        """创建超链接（简化版本）"""
        try:
            # 使用简化方法：直接添加带下划线的蓝色文本
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # 蓝色
            run.font.underline = True
            self._apply_font_to_run(run)  # 应用微软雅黑字体
            
            # 在括号中添加URL
            cleaned_url = self._clean_text(url)  # 清理URL文本
            url_run = paragraph.add_run(f" ({cleaned_url})")
            url_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)  # 灰色
            url_run.font.size = Pt(9)
            self._apply_font_to_run(url_run)  # 应用微软雅黑字体
            
            return run
        except Exception as e:
            # 如果超链接创建失败，添加普通文本
            self.logger.warning(f"超链接创建失败: {e}, 使用普通文本代替")
            cleaned_text = self._clean_text(text)  # 清理链接文本
            cleaned_url = self._clean_text(url)   # 清理URL文本
            run = paragraph.add_run(f"{cleaned_text} ({cleaned_url})")
            self._apply_font_to_run(run)  # 应用微软雅黑字体
            return run
    
    def _process_unordered_list(self, content: str, indent: int = 0):
        """处理无序列表"""
        cleaned_content = self._clean_text(content)  # 只使用基础清理
        paragraph = self.document.add_paragraph(cleaned_content, style='List Bullet')
        # 根据缩进调整列表级别
        if indent > 0:
            paragraph.paragraph_format.left_indent = Inches(indent * 0.25)
        # 确保列表项也使用微软雅黑字体
        for run in paragraph.runs:
            self._apply_font_to_run(run)
        self.logger.debug(f"添加无序列表项: {cleaned_content}")
    
    def _process_ordered_list(self, content: str, indent: int = 0):
        """处理有序列表"""
        cleaned_content = self._clean_text(content)  # 只使用基础清理
        paragraph = self.document.add_paragraph(cleaned_content, style='List Number')
        # 根据缩进调整列表级别
        if indent > 0:
            paragraph.paragraph_format.left_indent = Inches(indent * 0.25)
        # 确保列表项也使用微软雅黑字体
        for run in paragraph.runs:
            self._apply_font_to_run(run)
        self.logger.debug(f"添加有序列表项: {cleaned_content}")
    
    def _process_blockquote(self, content: str):
        """处理引用块"""
        cleaned_content = self._clean_text(content)  # 清理文本
        paragraph = self.document.add_paragraph(cleaned_content)
        try:
            paragraph.style = self.document.styles['Quote']
        except KeyError:
            # 如果样式不存在，手动设置格式
            paragraph.paragraph_format.left_indent = Inches(0.5)
            for run in paragraph.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x5a, 0x5a, 0x5a)
                self._apply_font_to_run(run)  # 应用微软雅黑字体
        self.logger.debug(f"添加引用: {cleaned_content}")
    
    def _process_horizontal_rule(self):
        """处理水平分割线"""
        paragraph = self.document.add_paragraph()
        rule_text = "─" * 50  # 使用更标准的水平线字符
        cleaned_rule = self._clean_text(rule_text)  # 清理分割线文本
        run = paragraph.add_run(cleaned_rule)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_font_to_run(run)  # 应用微软雅黑字体
        self.logger.debug("添加水平分割线")
    
    def _process_code_block(self, code_content: str, language: str = ""):
        """处理代码块"""
        cleaned_code = self._clean_text(code_content)  # 清理代码内容
        paragraph = self.document.add_paragraph(cleaned_code)
        try:
            paragraph.style = self.document.styles['CodeBlock']
        except KeyError:
            # 如果样式不存在，手动设置格式
            paragraph.paragraph_format.left_indent = Inches(0.5)
            for run in paragraph.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        
        if language:
            # 添加语言标识
            cleaned_language = self._clean_text(language)  # 清理语言标识
            lang_paragraph = self.document.add_paragraph(f"语言: {cleaned_language}")
            lang_paragraph.paragraph_format.space_after = Pt(0)
            for run in lang_paragraph.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                self._apply_font_to_run(run)  # 应用微软雅黑字体
        
        self.logger.debug(f"添加代码块: 语言={language}, 行数={len(code_content.split())}")
    
    def _process_table_block(self, table_lines: List[str]):
        """处理表格块 - 增强版本支持对齐和样式"""
        if len(table_lines) < 2:
            return
        
        # 解析表格数据和对齐信息
        table_data = self._parse_enhanced_table(table_lines)
        if not table_data:
            return
        
        # 创建和配置表格
        table = self._create_word_table(table_data)
        if not table:
            return
        
        # 填充表格内容
        self._populate_table_content(table, table_data)
        
        # 添加表格后的间距
        self._add_paragraph("")
        
        rows_count = len(table_data['rows'])
        cols_count = len(table_data['rows'][0]) if table_data['rows'] else 0
        self.logger.debug(f"添加增强表格: {rows_count}行 x {cols_count}列，对齐: {table_data['alignments']}")
    
    def _create_word_table(self, table_data: Dict[str, Any]):
        """创建Word表格并应用基础样式"""
        rows_data = table_data['rows']
        has_header = table_data['has_header']
        table_title = table_data.get('title', '')
        
        # 添加表格标题（如果存在）
        if table_title:
            self._add_table_title(table_title)
        
        # 创建Word表格
        if not rows_data or not rows_data[0]:
            return None
            
        table = self.document.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        
        # 应用表格样式
        self._apply_table_style(table, has_header)
        
        return table
    
    def _populate_table_content(self, table, table_data: Dict[str, Any]):
        """填充表格内容并应用格式"""
        rows_data = table_data['rows']
        alignment_info = table_data['alignments']
        has_header = table_data['has_header']
        
        for i, row in enumerate(rows_data):
            self._populate_table_row(table, i, row, alignment_info, has_header)
    
    def _populate_table_row(self, table, row_index: int, row_data: List[str], 
                           alignment_info: List[str], has_header: bool):
        """填充单行表格数据"""
        for j, cell_content in enumerate(row_data):
            if j >= len(table.rows[row_index].cells):
                continue
                
            cell = table.rows[row_index].cells[j]
            # 对表格内容进行温和清理，保留正常文字
            cleaned_content = self._clean_text(cell_content)  # 只使用基础清理
            cell.text = cleaned_content
            
            # 应用微软雅黑字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._apply_font_to_run(run)
            
            # 应用对齐方式
            if j < len(alignment_info):
                self._apply_cell_alignment(cell, alignment_info[j])
            
            # 如果是表头行，应用粗体
            if has_header and row_index == 0:
                self._apply_header_formatting(cell)
    
    def _apply_header_formatting(self, cell):
        """应用表头格式"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                self._apply_font_to_run(run)  # 确保表头也使用微软雅黑
    
    def _parse_enhanced_table(self, table_lines: List[str]) -> Optional[Dict[str, Any]]:
        """解析增强表格，提取数据和对齐信息"""
        if not table_lines:
            return None
        
        rows = []
        alignment_row_index = -1
        alignments = []
        
        # 查找表格行和对齐行
        for i, line in enumerate(table_lines):
            line = line.strip()
            if not line.startswith('|') or not line.endswith('|'):
                continue
                
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # 检查是否为对齐行 (包含 :---, ---:, :---: 等)
            if all(self._is_alignment_cell(cell) for cell in cells if cell):
                alignment_row_index = i
                alignments = [self._parse_alignment(cell) for cell in cells]
                continue
            
            # 跳过空行
            if cells and any(cell.strip() for cell in cells):
                rows.append(cells)
        
        if not rows:
            return None
        
        # 如果没有找到对齐行，使用默认左对齐
        if not alignments:
            alignments = ['left'] * len(rows[0])
        
        return {
            'rows': rows,
            'alignments': alignments,
            'has_header': alignment_row_index >= 0  # 有对齐行说明有表头
        }
    
    def _is_alignment_cell(self, cell: str) -> bool:
        """检查是否为对齐单元格"""
        cell = cell.strip()
        if not cell:
            return False
        return bool(re.match(r'^:?-+:?$', cell))
    
    def _parse_alignment(self, alignment_cell: str) -> str:
        """解析对齐方式"""
        cell = alignment_cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            return 'center'
        elif cell.endswith(':'):
            return 'right'
        else:
            return 'left'
    
    def _apply_table_style(self, table, has_header: bool):
        """应用表格样式"""
        # 设置基础样式
        table.style = 'Table Grid'
        
        # 设置表格属性
        table.autofit = True
        
        # 如果有表头，设置表头样式
        if has_header and len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # 设置表头背景色（浅灰色）
                shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _apply_cell_alignment(self, cell, alignment: str):
        """应用单元格对齐"""
        for paragraph in cell.paragraphs:
            if alignment == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:  # left or default
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_table_title(self, title: str):
        """添加表格标题"""
        cleaned_title = self._clean_text(title)  # 清理标题文本
        title_paragraph = self.document.add_paragraph(cleaned_title)
        try:
            title_paragraph.style = self.document.styles['TableTitle']
        except KeyError:
            # 如果样式不存在，手动设置格式
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                self._apply_font_to_run(run)  # 应用微软雅黑字体
        self.logger.debug(f"添加表格标题: {cleaned_title}")
    
    def _detect_table_title(self, text_before_table: str) -> str:
        """检测表格标题（从表格前的文本中）"""
        # 简单的标题检测：如果表格前一行是短文本且不包含特殊字符，视为标题
        if not text_before_table:
            return ""
        
        lines = text_before_table.strip().split('\n')
        if lines:
            last_line = lines[-1].strip()
            if len(last_line) < 50 and not any(char in last_line for char in ['#', '*', '`', '[']):
                return last_line
        return ""
    
    def _process_chart(self, chart_info: Dict[str, str]):
        """处理图表（集成图表渲染引擎）"""
        chart_type = chart_info['type']
        chart_code = chart_info['code']
        
        try:
            # 尝试渲染真实图表
            image_path = self._render_chart_to_image(chart_type, chart_code)
            
            if image_path and image_path.exists():
                # 成功渲染，插入图片
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中
                run = paragraph.add_run()
                
                # 添加图片到文档
                try:
                    # 智能计算图片尺寸
                    optimal_width = self._calculate_optimal_image_width(image_path, chart_type)
                    run.add_picture(str(image_path), width=optimal_width)
                    
                    self.logger.info(f"成功插入{chart_type}图表图片: {image_path}, 宽度: {optimal_width.inches:.1f}英寸")
                    
                    # 清理临时文件
                    try:
                        image_path.unlink()
                    except Exception as e:
                        self.logger.warning(f"清理临时图片文件失败: {e}")
                        
                except Exception as e:
                    self.logger.error(f"插入图片失败: {e}")
                    self._add_chart_fallback(chart_type, chart_code)
            else:
                # 渲染失败，使用降级方案
                self._add_chart_fallback(chart_type, chart_code)
                
        except Exception as e:
            self.logger.error(f"图表渲染异常: {e}")
            self._add_chart_fallback(chart_type, chart_code)
    
    def _render_chart_to_image(self, chart_type: str, chart_code: str) -> Optional[Any]:
        """渲染图表为图片文件
        
        Args:
            chart_type: 图表类型 (mermaid, plantuml)
            chart_code: 图表代码
            
        Returns:
            图片文件路径，失败时返回None
        """
        try:
            import tempfile
            from pathlib import Path
            
            # 创建临时输出目录
            temp_dir = Path(tempfile.gettempdir()) / "md2doc_charts"
            temp_dir.mkdir(exist_ok=True)
            
            # 生成唯一文件名
            import hashlib
            content_hash = hashlib.md5(chart_code.encode()).hexdigest()[:8]
            output_path = temp_dir / f"{chart_type}_{content_hash}.png"
            
            if chart_type.lower() == 'mermaid':
                return self._render_mermaid_chart(chart_code, output_path)
            elif chart_type.lower() == 'plantuml':
                return self._render_plantuml_chart(chart_code, output_path)
            else:
                self.logger.warning(f"不支持的图表类型: {chart_type}")
                return None
                
        except Exception as e:
            self.logger.error(f"图表渲染准备失败: {e}")
            return None
    
    def _render_mermaid_chart(self, chart_code: str, output_path: Any) -> Optional[Any]:
        """渲染Mermaid图表
        
        Args:
            chart_code: Mermaid代码
            output_path: 输出路径
            
        Returns:
            成功时返回文件路径，失败时返回None
        """
        try:
            from ..engines.chart_detector import ChartInfo, ChartType
            from ..engines.mermaid_engine import MermaidEngine
            
            # 创建图表信息对象
            chart_info = ChartInfo(
                chart_type=ChartType.MERMAID,
                content=chart_code,
                language="mermaid"
            )
            
            # 创建并配置Mermaid引擎
            engine = MermaidEngine()
            
            # 渲染图表
            result = engine.render(chart_info, output_path)
            
            if isinstance(result, bytes):
                # 结果是字节数据，保存到文件
                with open(output_path, 'wb') as f:
                    f.write(result)
                return output_path
            elif hasattr(result, 'exists') and result.exists():
                # 结果是文件路径
                return result
            else:
                self.logger.warning("Mermaid渲染返回了无效结果")
                return None
                
        except Exception as e:
            self.logger.error(f"Mermaid图表渲染失败: {e}")
            return None
    
    def _render_plantuml_chart(self, chart_code: str, output_path: Any) -> Optional[Any]:
        """渲染PlantUML图表
        
        Args:
            chart_code: PlantUML代码
            output_path: 输出路径
            
        Returns:
            成功时返回文件路径，失败时返回None
        """
        try:
            from ..engines.chart_detector import ChartInfo, ChartType
            from ..engines.plantuml_engine import PlantUMLEngine
            
            # 创建图表信息对象
            chart_info = ChartInfo(
                chart_type=ChartType.PLANTUML,
                content=chart_code,
                language="plantuml"
            )
            
            # 创建PlantUML引擎
            engine = PlantUMLEngine()
            
            # 渲染图表
            success, result_path, error = engine.render(chart_info, output_path)
            
            if success and result_path and result_path.exists():
                return result_path
            else:
                self.logger.warning(f"PlantUML渲染失败: {error}")
                return None
                
        except Exception as e:
            self.logger.error(f"PlantUML图表渲染失败: {e}")
            return None
    
    def _add_chart_fallback(self, chart_type: str, chart_code: str):
        """添加图表降级方案（显示代码块）
        
        Args:
            chart_type: 图表类型
            chart_code: 图表代码
        """
        # 添加图表标题
        paragraph = self.document.add_paragraph()
        title_text = f"[{chart_type.upper()}图表]"
        cleaned_title = self._clean_text(title_text)  # 清理标题文本
        run = paragraph.add_run(cleaned_title)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        self._apply_font_to_run(run)  # 应用微软雅黑字体
        
        # 添加图表代码
        cleaned_code = self._clean_text(chart_code)  # 清理代码文本
        code_paragraph = self.document.add_paragraph(cleaned_code)
        try:
            code_paragraph.style = self.document.styles['CodeBlock']
        except KeyError:
            for run in code_paragraph.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                self._apply_font_to_run(run)  # 应用微软雅黑字体
        
        self.logger.info(f"使用降级方案显示{chart_type}图表代码")
    
    def _calculate_optimal_image_width(self, image_path: Any, chart_type: str) -> Any:
        """计算图片的显示宽度 - 1:1比例，不缩放
        
        Args:
            image_path: 图片文件路径
            chart_type: 图表类型
            
        Returns:
            图片原始物理宽度（Inches对象）
        """
        try:
            from PIL import Image
            from docx.shared import Inches
            
            # 读取图片尺寸和DPI信息
            with Image.open(image_path) as img:
                width, height = img.size
                
                # 获取DPI信息，默认96DPI
                dpi_x, dpi_y = img.info.get('dpi', (96, 96))
                if isinstance(dpi_x, tuple):
                    dpi_x = dpi_x[0] if dpi_x[0] > 0 else 96
                if isinstance(dpi_y, tuple):
                    dpi_y = dpi_y[0] if dpi_y[0] > 0 else 96
                    
                # 计算图片的实际物理尺寸（英寸）
                physical_width_inches = width / dpi_x
                physical_height_inches = height / dpi_y
                
                self.logger.info(f"图片信息: {width}x{height}px, DPI: {dpi_x}x{dpi_y}, 物理尺寸: {physical_width_inches:.2f}x{physical_height_inches:.2f}英寸")
                
                # 1:1比例显示，完全按照原始物理尺寸
                optimal_width = Inches(physical_width_inches)
                
                self.logger.info(f"使用1:1比例显示: {optimal_width.inches:.2f}英寸")
                return optimal_width
                
        except ImportError:
            self.logger.warning("PIL库未安装，使用默认图片尺寸")
        except Exception as e:
            self.logger.warning(f"计算图片尺寸失败: {e}")
        
        # 降级方案：根据图表类型返回固定的合理尺寸
        from docx.shared import Inches
        if chart_type.lower() == 'mermaid':
            return Inches(4.0)  # Mermaid适中尺寸
        elif chart_type.lower() == 'plantuml':
            return Inches(3.5)  # PlantUML较小尺寸（通常比较紧凑）
        elif chart_type.lower() == 'image':
            return Inches(4.5)  # 普通图片默认尺寸
        else:
            return Inches(4.0)  # 其他类型默认
    
    def _add_paragraph(self, text: str = "") -> Any:
        """添加段落的辅助方法"""
        cleaned_text = self._clean_text(text) if text else text
        paragraph = self.document.add_paragraph(cleaned_text)
        # 确保段落中的所有运行都使用正确的字体
        for run in paragraph.runs:
            self._apply_font_to_run(run)
        return paragraph
    
    def save_document(self, file_path: str):
        """保存Word文档
        
        Args:
            file_path: 保存路径
        """
        try:
            self.document.save(file_path)
            self.logger.info(f"文档保存成功: {file_path}")
        except Exception as e:
            self.logger.error(f"文档保存失败: {e}")
            raise


# 格式转换工具函数
def convert_markdown_to_docx(markdown_file: str, output_file: str) -> bool:
    """转换Markdown文件到Word文档
    
    Args:
        markdown_file: Markdown文件路径
        output_file: 输出Word文件路径
        
    Returns:
        转换是否成功
    """
    try:
        # 读取Markdown文件
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # 创建格式转换器
        converter = FormatConverter()
        
        # 执行转换
        converter.convert_markdown_to_word(markdown_content)
        
        # 保存文档
        converter.save_document(output_file)
        
        return True
        
    except Exception as e:
        logging.error(f"格式转换失败: {e}")
        return False


if __name__ == "__main__":
    # 测试代码
    logging.basicConfig(level=logging.DEBUG)
    
    test_markdown = """
# 测试文档

这是一个**粗体**和*斜体*的测试。

## 代码示例

这里有一些`内联代码`。

```python
def hello_world():
    print("Hello, World!")
```

## 列表测试

- 无序列表项1
- 无序列表项2
  - 嵌套项1
  - 嵌套项2

1. 有序列表项1
2. 有序列表项2

## 链接和图片

这是一个[链接](https://www.example.com)。

![测试图片](test.png)

## 引用

> 这是一个引用块。
> 可以有多行。

## 表格

| 列1 | 列2 | 列3 |
|-----|-----|-----|
| 数据1 | 数据2 | 数据3 |
| 数据4 | 数据5 | 数据6 |

---

## 图表测试

```mermaid
graph TD
    A[开始] --> B[处理]
    B --> C[结束]
```
"""
    
    converter = FormatConverter()
    document = converter.convert_markdown_to_word(test_markdown)
    converter.save_document("test_output.docx")
    print("测试转换完成！")
