#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
优化的格式转换器模块 - 修复换行问题
特别针对代码块和命令行内容的换行处理进行优化
"""

import re
import logging
import tempfile
import hashlib
from typing import Dict, List, Any, Optional, Tuple, TYPE_CHECKING
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

if TYPE_CHECKING:
    from docx import Document as DocxDocument
    from docx.document import Document as DocxDocumentType
else:
    DocxDocument = Any
    DocxDocumentType = Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
except ImportError as e:
    logging.error(f"python-docx库未安装: {e}")
    raise ImportError("请安装python-docx库: pip install python-docx")


class ContentType(Enum):
    """内容类型枚举"""
    NORMAL_TEXT = "normal_text"
    CODE_BLOCK = "code_block"
    INLINE_CODE = "inline_code"
    COMMAND_LINE = "command_line"
    TABLE_CELL = "table_cell"


class OptimizedFormatConverter:
    """优化的格式转换器 - 修复换行问题"""
    
    def __init__(self, document: Optional[Any] = None):
        """初始化优化格式转换器"""
        self.document = document or Document()
        self.logger = logging.getLogger(__name__)
        
        # 立即设置文档字体
        self._force_set_document_font()
        
        # 初始化样式
        self._init_word_styles()
        
    def _force_set_document_font(self):
        """强制设置文档字体为微软雅黑"""
        try:
            # 设置Normal样式的字体
            normal_style = self.document.styles['Normal']
            normal_style.font.name = 'Microsoft YaHei'
            normal_style.font.size = Pt(12)
            
            # 设置所有标题样式的字体
            for level in range(1, 10):
                try:
                    heading_style = self.document.styles[f'Heading {level}']
                    heading_style.font.name = 'Microsoft YaHei'
                except KeyError:
                    pass
            
            self.logger.info("文档字体设置完成")
            
        except Exception as e:
            self.logger.error(f"字体设置失败: {e}")
    
    def _init_word_styles(self):
        """初始化Word文档样式"""
        styles = self.document.styles
        
        # 创建代码块样式
        try:
            code_block_style = styles.add_style('OptimizedCodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_block_style.base_style = styles['Normal']
            code_block_font = code_block_style.font
            code_block_font.name = 'Consolas'
            code_block_font.size = Pt(10)
            code_block_font.color.rgb = RGBColor(0x1f, 0x1f, 0x1f)
            
            # 设置段落格式
            paragraph_format = code_block_style.paragraph_format
            paragraph_format.left_indent = Inches(0.5)
            paragraph_format.space_after = Pt(6)
            paragraph_format.space_before = Pt(6)
            
        except ValueError:
            pass  # 样式可能已存在
            
        # 创建内联代码样式
        try:
            inline_code_style = styles.add_style('OptimizedInlineCode', WD_STYLE_TYPE.CHARACTER)
            inline_code_font = inline_code_style.font
            inline_code_font.name = 'Consolas'
            inline_code_font.size = Pt(10)
            inline_code_font.color.rgb = RGBColor(0x1f, 0x1f, 0x1f)
        except ValueError:
            pass
    
    def _clean_text_optimized(self, text: str, content_type: ContentType = ContentType.NORMAL_TEXT) -> str:
        """优化的文本清理函数 - 根据内容类型智能处理换行
        
        Args:
            text: 需要清理的文本
            content_type: 内容类型
            
        Returns:
            清理后的文本
        """
        if not text:
            return text
        
        cleaned_text = text
        
        # 根据内容类型采用不同的清理策略
        if content_type == ContentType.CODE_BLOCK:
            # 代码块：保留所有换行符，只清理有害字符
            cleaned_text = self._clean_code_block_text(cleaned_text)
        elif content_type == ContentType.COMMAND_LINE:
            # 命令行：保留换行符，但清理多余空格
            cleaned_text = self._clean_command_line_text(cleaned_text)
        elif content_type == ContentType.INLINE_CODE:
            # 内联代码：保留内部空格，移除换行
            cleaned_text = self._clean_inline_code_text(cleaned_text)
        elif content_type == ContentType.TABLE_CELL:
            # 表格单元格：温和清理，保留基本格式
            cleaned_text = self._clean_table_cell_text(cleaned_text)
        else:
            # 普通文本：标准清理，将换行替换为空格
            cleaned_text = self._clean_normal_text(cleaned_text)
        
        return cleaned_text
    
    def _clean_code_block_text(self, text: str) -> str:
        """清理代码块文本 - 保留换行符"""
        if not text:
            return text
        
        # 只移除有害的控制字符，保留换行符和制表符
        cleaned_text = text
        
        # 移除零宽字符
        cleaned_text = re.sub(r'[\u200B-\u200F]', '', cleaned_text)
        cleaned_text = re.sub(r'[\uFEFF]', '', cleaned_text)
        cleaned_text = re.sub(r'[\uFE0E\uFE0F]', '', cleaned_text)
        
        # 移除有害控制字符，但保留 \\n、\\r、\\t
        cleaned_text = re.sub(r'[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F]', '', cleaned_text)
        
        # 移除软连字符等问题字符
        cleaned_text = re.sub(r'[\u00AD]', '', cleaned_text)
        
        # 标准化行尾：统一使用 \n
        cleaned_text = re.sub(r'\r\n|\r', '\n', cleaned_text)
        
        # 移除首尾空格，但保留内部结构
        cleaned_text = cleaned_text.rstrip()
        
        return cleaned_text
    
    def _clean_command_line_text(self, text: str) -> str:
        """清理命令行文本 - 保留换行，清理多余空格"""
        if not text:
            return text
        
        cleaned_text = self._clean_code_block_text(text)  # 基础清理
        
        # 清理每行的首尾空格，但保留换行
        lines = cleaned_text.split('\n')
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        cleaned_text = '\n'.join(cleaned_lines)
        
        return cleaned_text
    
    def _clean_inline_code_text(self, text: str) -> str:
        """清理内联代码文本 - 移除换行，保留内部空格"""
        if not text:
            return text
        
        # 基础清理
        cleaned_text = text
        cleaned_text = re.sub(r'[\u200B-\u200F]', '', cleaned_text)
        cleaned_text = re.sub(r'[\uFEFF]', '', cleaned_text)
        
        # 将换行替换为空格
        cleaned_text = re.sub(r'[\r\n]+', ' ', cleaned_text)
        
        # 清理多余空格
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        cleaned_text = cleaned_text.strip()
        
        return cleaned_text
    
    def _clean_table_cell_text(self, text: str) -> str:
        """清理表格单元格文本 - 温和清理，包含全面的Markdown语法清理"""
        if not text:
            return text
        
        cleaned_text = text
        
        # 清理所有常见的Markdown语法标记 - 重要修复！
        cleaned_text = re.sub(r'\*\*\*([^*]+)\*\*\*', r'\1', cleaned_text)  # ***粗斜体***
        cleaned_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned_text)      # **粗体**
        cleaned_text = re.sub(r'\*([^*\n]+)\*', r'\1', cleaned_text)         # *斜体*
        cleaned_text = re.sub(r'___([^_]+)___', r'\1', cleaned_text)         # ___粗斜体___
        cleaned_text = re.sub(r'__([^_]+)__', r'\1', cleaned_text)           # __粗体__
        cleaned_text = re.sub(r'_([^_\n]+)_', r'\1', cleaned_text)           # _斜体_
        cleaned_text = re.sub(r'`([^`]+)`', r'\1', cleaned_text)             # `代码`
        cleaned_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned_text) # [文本](链接)
        cleaned_text = re.sub(r'#{1,6}\s+', '', cleaned_text)               # ### 标题标记
        cleaned_text = re.sub(r'~~([^~]+)~~', r'\1', cleaned_text)          # ~~删除线~~
        cleaned_text = re.sub(r'\+\+([^+]+)\+\+', r'\1', cleaned_text)      # ++插入++
        cleaned_text = re.sub(r'==([^=]+)==', r'\1', cleaned_text)          # ==高亮==
        
        # 清理列表标记
        cleaned_text = re.sub(r'^\s*[-*+]\s+', '', cleaned_text)            # 无序列表
        cleaned_text = re.sub(r'^\s*\d+\.\s+', '', cleaned_text)            # 有序列表
        
        # 清理引用标记
        cleaned_text = re.sub(r'^\s*>\s*', '', cleaned_text)                # > 引用
        
        # 基础清理
        cleaned_text = re.sub(r'[\u200B-\u200F]', '', cleaned_text)
        cleaned_text = re.sub(r'[\uFEFF]', '', cleaned_text)
        
        # 将换行替换为空格
        cleaned_text = re.sub(r'[\r\n]+', ' ', cleaned_text)
        
        # 清理多余空格
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        cleaned_text = cleaned_text.strip()
        
        return cleaned_text
    
    def _clean_normal_text(self, text: str) -> str:
        """清理普通文本 - 标准清理"""
        if not text:
            return text
        
        cleaned_text = text
        
        # 清理Markdown语法标记
        cleaned_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned_text)  # **粗体**
        cleaned_text = re.sub(r'\*([^*\n]+)\*', r'\1', cleaned_text)      # *斜体*
        cleaned_text = re.sub(r'`([^`]+)`', r'\1', cleaned_text)           # `代码`
        cleaned_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned_text)  # [文本](链接)
        
        # 移除有害字符
        cleaned_text = re.sub(r'[\u200B-\u200F]', '', cleaned_text)
        cleaned_text = re.sub(r'[\uFEFF]', '', cleaned_text)
        cleaned_text = re.sub(r'[\u00AD]', '', cleaned_text)
        
        # 标准化空白字符（将换行替换为空格）
        cleaned_text = re.sub(r'[\s]+', ' ', cleaned_text)
        cleaned_text = cleaned_text.strip()
        
        return cleaned_text
    
    def _apply_font_to_run(self, run, font_name: str = 'Microsoft YaHei'):
        """为run对象应用指定字体"""
        try:
            run.font.name = font_name
            
            # XML级别设置
            from docx.oxml.ns import qn
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.find(qn('w:rFonts'))
            if r_fonts is None:
                r_fonts = parse_xml(f'<w:rFonts {r_pr.nsmap} />')
                r_pr.append(r_fonts)
            
            r_fonts.set(qn('w:ascii'), font_name)
            r_fonts.set(qn('w:hAnsi'), font_name)
            r_fonts.set(qn('w:eastAsia'), font_name)
            r_fonts.set(qn('w:cs'), font_name)
            
        except Exception as e:
            self.logger.warning(f"字体设置失败: {e}")
            try:
                run.font.name = font_name
            except Exception:
                pass
    
    def convert_markdown_to_word(self, markdown_text: str) -> Any:
        """优化版本的Markdown到Word转换"""
        self.logger.info("开始优化版本的Markdown到Word格式转换")
        
        # 预处理：分离图表代码块
        processed_text, chart_blocks = self._extract_chart_blocks(markdown_text)
        
        # 按行处理
        lines = processed_text.split('\n')
        current_table = []
        in_code_block = False
        code_block_content = []
        code_language = ""
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 处理代码块
            if line.strip().startswith('```'):
                if not in_code_block:
                    # 开始代码块
                    in_code_block = True
                    code_language = line.strip()[3:]
                    code_block_content = []
                else:
                    # 结束代码块 - 使用优化处理
                    self._process_code_block_optimized('\n'.join(code_block_content), code_language)
                    in_code_block = False
                    code_block_content = []
                    code_language = ""
                i += 1
                continue
                
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
            
            # 处理表格
            if self._is_table_line(line):
                current_table.append(line)
                i += 1
                continue
            elif current_table:
                # 表格结束
                self._process_table_block_optimized(current_table)
                current_table = []
            
            # 处理普通行
            self._process_line_optimized(line, chart_blocks)
            i += 1
        
        # 处理剩余的表格
        if current_table:
            self._process_table_block_optimized(current_table)
            
        self.logger.info("优化版本Markdown到Word格式转换完成")
        return self.document
    
    def _process_code_block_optimized(self, code_content: str, language: str = ""):
        """优化的代码块处理 - 完美保留换行"""
        if not code_content:
            return
        
        # 使用优化清理，保留换行符
        cleaned_code = self._clean_text_optimized(code_content, ContentType.CODE_BLOCK)
        
        # 如果有语言标识，添加语言标签
        if language:
            lang_paragraph = self.document.add_paragraph()
            lang_run = lang_paragraph.add_run(f"语言: {language}")
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            self._apply_font_to_run(lang_run)
            lang_paragraph.paragraph_format.space_after = Pt(0)
        
        # 按行处理代码内容，每行单独一个段落以保持换行
        code_lines = cleaned_code.split('\n')
        for line in code_lines:
            code_paragraph = self.document.add_paragraph()
            
            # 设置代码块段落格式
            code_paragraph.paragraph_format.left_indent = Inches(0.5)
            code_paragraph.paragraph_format.space_after = Pt(0)
            code_paragraph.paragraph_format.space_before = Pt(0)
            
            # 添加代码内容
            if line.strip():  # 非空行
                code_run = code_paragraph.add_run(line)
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(10)
                code_run.font.color.rgb = RGBColor(0x1f, 0x1f, 0x1f)
            else:  # 空行
                code_run = code_paragraph.add_run("")  # 空run保持行距
        
        # 代码块后添加间距
        self.document.add_paragraph("")
        
        self.logger.debug(f"添加优化代码块: 语言={language}, 行数={len(code_lines)}")
    
    def _process_table_block_optimized(self, table_lines: List[str]):
        """优化的表格处理"""
        if len(table_lines) < 2:
            return
        
        # 解析表格数据
        table_data = self._parse_enhanced_table_optimized(table_lines)
        if not table_data:
            return
        
        # 创建Word表格
        table = self._create_word_table_optimized(table_data)
        if not table:
            return
        
        # 填充表格内容（使用优化清理）
        self._populate_table_content_optimized(table, table_data)
        
        # 添加表格后的间距
        self.document.add_paragraph("")
        
        rows_count = len(table_data['rows'])
        cols_count = len(table_data['rows'][0]) if table_data['rows'] else 0
        self.logger.debug(f"添加优化表格: {rows_count}行 x {cols_count}列")
    
    def _parse_enhanced_table_optimized(self, table_lines: List[str]) -> Optional[Dict[str, Any]]:
        """优化的表格解析"""
        if not table_lines:
            return None
        
        rows = []
        alignment_row_index = -1
        alignments = []
        
        for i, line in enumerate(table_lines):
            line = line.strip()
            if not line.startswith('|') or not line.endswith('|'):
                continue
                
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # 检查是否为对齐行
            if all(self._is_alignment_cell(cell) for cell in cells if cell):
                alignment_row_index = i
                alignments = [self._parse_alignment(cell) for cell in cells]
                continue
            
            if cells and any(cell.strip() for cell in cells):
                rows.append(cells)
        
        if not rows:
            return None
        
        if not alignments:
            alignments = ['left'] * len(rows[0])
        
        return {
            'rows': rows,
            'alignments': alignments,
            'has_header': alignment_row_index >= 0
        }
    
    def _create_word_table_optimized(self, table_data: Dict[str, Any]):
        """创建优化的Word表格"""
        rows_data = table_data['rows']
        
        if not rows_data or not rows_data[0]:
            return None
            
        table = self.document.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        table.style = 'Table Grid'
        table.autofit = True
        
        return table
    
    def _populate_table_content_optimized(self, table, table_data: Dict[str, Any]):
        """使用优化清理填充表格内容"""
        rows_data = table_data['rows']
        alignment_info = table_data['alignments']
        has_header = table_data['has_header']
        
        for i, row in enumerate(rows_data):
            for j, cell_content in enumerate(row):
                if j >= len(table.rows[i].cells):
                    continue
                    
                cell = table.rows[i].cells[j]
                
                # 使用优化清理处理表格单元格内容
                cleaned_content = self._clean_text_optimized(cell_content, ContentType.TABLE_CELL)
                cell.text = cleaned_content
                
                # 应用字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._apply_font_to_run(run)
                
                # 应用对齐
                if j < len(alignment_info):
                    self._apply_cell_alignment(cell, alignment_info[j])
                
                # 表头格式
                if has_header and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            self._apply_font_to_run(run)
    
    def _process_line_optimized(self, line: str, chart_blocks: Dict[str, Dict[str, str]]):
        """优化的行处理"""
        line = line.strip()
        if not line:
            self.document.add_paragraph("")
            return
            
        # 处理标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            self._process_heading_optimized(title, level)
            return
        
        # 处理列表
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            content = list_match.group(3)
            if marker.isdigit() or marker.endswith('.'):
                self._process_ordered_list_optimized(content, indent)
            else:
                self._process_unordered_list_optimized(content, indent)
            return
        
        # 普通段落
        cleaned_text = self._clean_text_optimized(line, ContentType.NORMAL_TEXT)
        paragraph = self.document.add_paragraph()
        self._process_paragraph_with_formatting(paragraph, cleaned_text)
    
    def _process_heading_optimized(self, title: str, level: int):
        """优化的标题处理"""
        cleaned_title = self._clean_text_optimized(title, ContentType.NORMAL_TEXT)
        heading_style = f'Heading {min(level, 6)}'
        paragraph = self.document.add_paragraph(cleaned_title, style=heading_style)
        for run in paragraph.runs:
            self._apply_font_to_run(run)
        self.logger.debug(f"添加标题: 级别={level}, 内容={cleaned_title}")
    
    def _process_unordered_list_optimized(self, content: str, indent: int = 0):
        """优化的无序列表处理"""
        cleaned_content = self._clean_text_optimized(content, ContentType.NORMAL_TEXT)
        paragraph = self.document.add_paragraph(cleaned_content, style='List Bullet')
        if indent > 0:
            paragraph.paragraph_format.left_indent = Inches(indent * 0.25)
        for run in paragraph.runs:
            self._apply_font_to_run(run)
    
    def _process_ordered_list_optimized(self, content: str, indent: int = 0):
        """优化的有序列表处理"""
        cleaned_content = self._clean_text_optimized(content, ContentType.NORMAL_TEXT)
        paragraph = self.document.add_paragraph(cleaned_content, style='List Number')
        if indent > 0:
            paragraph.paragraph_format.left_indent = Inches(indent * 0.25)
        for run in paragraph.runs:
            self._apply_font_to_run(run)
    
    def _process_paragraph_with_formatting(self, paragraph, text: str):
        """处理段落格式（粗体、斜体等）"""
        # 简化处理：直接添加文本
        run = paragraph.add_run(text)
        self._apply_font_to_run(run)
    
    # 辅助方法
    def _extract_chart_blocks(self, text: str) -> Tuple[str, Dict[str, Dict[str, str]]]:
        """提取图表代码块"""
        return text, {}  # 简化版本
    
    def _is_table_line(self, line: str) -> bool:
        """判断是否为表格行"""
        stripped = line.strip()
        return stripped.startswith('|') and stripped.endswith('|')
    
    def _is_alignment_cell(self, cell: str) -> bool:
        """检查是否为对齐单元格"""
        cell = cell.strip()
        if not cell:
            return False
        return bool(re.match(r'^:?-+:?$', cell))
    
    def _parse_alignment(self, alignment_cell: str) -> str:
        """解析对齐方式"""
        cell = alignment_cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            return 'center'
        elif cell.endswith(':'):
            return 'right'
        else:
            return 'left'
    
    def _apply_cell_alignment(self, cell, alignment: str):
        """应用单元格对齐"""
        for paragraph in cell.paragraphs:
            if alignment == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def save_document(self, file_path: str):
        """保存Word文档"""
        try:
            self.document.save(file_path)
            self.logger.info(f"优化版本文档保存成功: {file_path}")
        except Exception as e:
            self.logger.error(f"文档保存失败: {e}")
            raise


if __name__ == "__main__":
    # 测试优化转换器
    logging.basicConfig(level=logging.INFO)
    
    test_markdown = """# 测试文档

## 代码块测试

这里有一些命令行示例：

```bash
# 检查SSH配置
ssh -T git@github.com

# 重新配置SSH密钥
ssh-keygen -t ed25519 -C "your-email@example.com"

# 将~/.ssh/id_ed25519.pub内容添加到GitHub
```

```python
def hello_world():
    print("Hello, World!")
    return True
```

## 表格测试

| 命令 | 说明 | 示例 |
|------|------|------|
| ssh -T | 测试连接 | ssh -T git@github.com |
| ssh-keygen | 生成密钥 | ssh-keygen -t ed25519 |

正常文本应该正确显示。
"""
    
    converter = OptimizedFormatConverter()
    document = converter.convert_markdown_to_word(test_markdown)
    converter.save_document("test_optimized_output.docx")
    print("优化版本测试转换完成！")
