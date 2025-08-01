"""
Word文档生成器

提供Word文档的生成功能：
- 基础文档结构生成
- 样式和格式应用
- 标题层级处理
- 段落和列表生成
"""

from typing import Dict, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
import logging

from .parser import ParseResult, MarkdownElement, HeadingElement, ParagraphElement, ListElement, CodeBlockElement, ChartElement, TableElement
from .config import ConfigManager
from ..engines.mermaid_engine import MermaidEngine
from ..engines.plantuml_engine import PlantUMLEngine


class WordDocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self, config: Optional[ConfigManager] = None):
        """初始化生成器
        
        Args:
            config: 配置管理器实例
        """
        self.config = config or ConfigManager()
        self.logger = logging.getLogger(__name__)
        self.document = None
        
        # 初始化图表渲染引擎
        self.mermaid_engine = MermaidEngine(self.config.get('chart', {}))
        self.plantuml_engine = PlantUMLEngine()
        
        # 图表缓存目录
        self.chart_cache_dir = Path(self.config.get('chart.cache_dir', 'temp/charts'))
        self.chart_cache_dir.mkdir(parents=True, exist_ok=True)
        
    def create_document(self):
        """创建新的Word文档
        
        Returns:
            Word文档对象
        """
        self.document = Document()
        self._setup_document_styles()
        self._setup_document_margins()
        return self.document
    
    def _setup_document_styles(self):
        """设置文档样式"""
        if not self.document:
            return
            
        styles = self.document.styles
        
        # 配置标题样式
        for level in range(1, 7):
            style_name = f'Heading {level}'
            if style_name in styles:
                heading_style = styles[style_name]
            else:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            
            # 设置字体
            font_name = self.config.get('document.font_name', '微软雅黑')
            heading_style.font.name = font_name
            heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
            # 设置字体大小
            base_size = self.config.get('document.font_size', 12)
            heading_sizes = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11}
            font_size = heading_sizes.get(level, base_size)
            if font_size:
                heading_style.font.size = Pt(font_size)
            heading_style.font.bold = True
            
            # 设置段落格式
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        
        # 配置正文样式
        normal_style = styles['Normal']
        font_name = self.config.get('document.font_name', '微软雅黑')
        normal_style.font.name = font_name
        normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        normal_style.font.size = Pt(self.config.get('document.font_size', 12))
        
        # 设置行间距
        line_spacing = self.config.get('document.line_spacing', 1.15)
        normal_style.paragraph_format.line_spacing = line_spacing
        
        # 创建代码样式
        if 'Code' not in styles:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(10)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
            
            # 设置背景色（浅灰色）
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), 'F8F8F8')
            code_style.element.pPr.append(shading_elm)
    
    def _setup_document_margins(self):
        """设置文档页边距"""
        if not self.document:
            return
            
        sections = self.document.sections
        for section in sections:
            margin_config = self.config.get('document.margin', {})
            section.top_margin = Inches(margin_config.get('top', 2.54))
            section.bottom_margin = Inches(margin_config.get('bottom', 2.54))
            section.left_margin = Inches(margin_config.get('left', 3.17))
            section.right_margin = Inches(margin_config.get('right', 3.17))
    
    def generate_from_parse_result(self, parse_result: ParseResult, output_path: Optional[Path] = None):
        """从解析结果生成Word文档
        
        Args:
            parse_result: Markdown解析结果
            output_path: 输出文件路径
            
        Returns:
            生成的Word文档
        """
        self.logger.info("开始生成Word文档")
        
        # 创建文档
        doc = self.create_document()
        
        # 处理每个元素
        for element in parse_result.elements:
            self._add_element_to_document(element)
        
        # 保存文档
        if output_path:
            self.save_document(doc, output_path)
            self.logger.info(f"文档已保存到: {output_path}")
        
        return doc
    
    def _add_element_to_document(self, element: MarkdownElement):
        """将解析元素添加到文档中
        
        Args:
            element: Markdown元素
        """
        if isinstance(element, HeadingElement):
            self._add_heading(element)
        elif isinstance(element, ParagraphElement):
            self._add_paragraph(element)
        elif isinstance(element, ListElement):
            self._add_list(element)
        elif isinstance(element, CodeBlockElement):
            self._add_code_block(element)
        elif isinstance(element, TableElement):
            self._add_table(element)
        elif isinstance(element, ChartElement):
            self._add_chart_placeholder(element)
        else:
            self.logger.warning(f"未知元素类型: {type(element)}")
    
    def _add_heading(self, heading: HeadingElement):
        """添加标题
        
        Args:
            heading: 标题元素
        """
        self.document.add_heading(heading.content, level=heading.level)
        self.logger.debug(f"添加标题: {heading.content} (级别 {heading.level})")
    
    def _add_paragraph(self, paragraph_element: ParagraphElement):
        """添加段落
        
        Args:
            paragraph_element: 段落元素
        """
        self.document.add_paragraph(paragraph_element.content)
        self.logger.debug(f"添加段落: {paragraph_element.content[:50]}...")
    
    def _add_list(self, list_element: ListElement):
        """添加列表
        
        Args:
            list_element: 列表元素
        """
        self.logger.debug(f"添加列表: {list_element.list_type}, {len(list_element.items)}项")
        
        for item in list_element.items:
            paragraph = self.document.add_paragraph()
            
            if list_element.list_type == "ordered":
                # 有序列表
                paragraph.style = self.document.styles['List Number']
            else:
                # 无序列表
                paragraph.style = self.document.styles['List Bullet']
            
            paragraph.add_run(item)
    
    def _add_code_block(self, code_element: CodeBlockElement):
        """添加代码块
        
        Args:
            code_element: 代码块元素
        """
        self.logger.debug(f"添加代码块: {code_element.language}")
        
        # 添加语言标识（如果有）
        if code_element.language:
            lang_paragraph = self.document.add_paragraph(f"代码 ({code_element.language}):")
            lang_paragraph.runs[0].font.bold = True
        
        # 添加代码内容
        code_paragraph = self.document.add_paragraph(code_element.content)
        if 'Code' in self.document.styles:
            code_paragraph.style = 'Code'
    
    def _add_table(self, table_element: TableElement):
        """添加表格
        
        Args:
            table_element: 表格元素
        """
        self.logger.debug(f"添加表格: {len(table_element.headers)}列, {len(table_element.rows)}行")
        
        # 创建表格
        table = self.document.add_table(
            rows=1 + len(table_element.rows), 
            cols=len(table_element.headers)
        )
        table.style = 'Table Grid'
        
        # 添加表头
        header_row = table.rows[0]
        for i, header in enumerate(table_element.headers):
            cell = header_row.cells[i]
            cell.text = header
            # 设置表头样式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 添加数据
        for row_index, row_data in enumerate(table_element.rows):
            table_row = table.rows[row_index + 1]
            for col_index, cell_data in enumerate(row_data):
                if col_index < len(table_row.cells):
                    table_row.cells[col_index].text = cell_data
    
    def _add_chart_placeholder(self, chart_element: ChartElement):
        """添加图表（渲染或占位符）
        
        Args:
            chart_element: 图表元素
        """
        self.logger.debug(f"处理图表: {chart_element.chart_type}")
        
        # 尝试渲染图表
        try:
            if chart_element.chart_type == "mermaid":
                self._render_mermaid_chart(chart_element)
            elif chart_element.chart_type == "plantuml":
                self._render_plantuml_chart(chart_element)
            else:
                # 其他类型的图表暂时显示占位符
                self._add_chart_placeholder_fallback(chart_element)
        except Exception as e:
            self.logger.warning(f"图表渲染失败: {e}")
            self._add_chart_placeholder_fallback(chart_element)
    
    def _render_mermaid_chart(self, chart_element: ChartElement):
        """渲染Mermaid图表
        
        Args:
            chart_element: 图表元素
        """
        from ..engines.chart_detector import ChartInfo, ChartType
        
        # 创建图表信息对象
        chart_info = ChartInfo(
            chart_type=ChartType.MERMAID,
            content=chart_element.content,
            language="mermaid",
            config=chart_element.attributes
        )
        
        # 检查是否能渲染
        if not self.mermaid_engine.can_render(chart_info):
            self.logger.warning("Mermaid引擎无法渲染此图表")
            self._add_chart_placeholder_fallback(chart_element)
            return
        
        # 生成唯一的文件名
        import hashlib
        content_hash = hashlib.md5(chart_element.content.encode()).hexdigest()[:8]
        chart_filename = f"mermaid_{content_hash}.png"
        chart_path = self.chart_cache_dir / chart_filename
        
        try:
            # 渲染图表
            self.logger.debug(f"渲染Mermaid图表到: {chart_path}")
            result = self.mermaid_engine.render(chart_info, chart_path)
            
            if isinstance(result, Path) and result.exists():
                # 成功渲染，插入图片到文档
                self._insert_chart_image(chart_path, chart_element)
            else:
                # 渲染失败，使用占位符
                self.logger.warning("图表渲染结果无效")
                self._add_chart_placeholder_fallback(chart_element)
                
        except Exception as e:
            self.logger.error(f"Mermaid图表渲染失败: {e}")
            self._add_chart_placeholder_fallback(chart_element)
    
    def _render_plantuml_chart(self, chart_element: ChartElement):
        """渲染PlantUML图表
        
        Args:
            chart_element: 图表元素
        """
        from ..engines.chart_detector import ChartInfo, ChartType
        
        # 创建图表信息对象
        chart_info = ChartInfo(
            chart_type=ChartType.PLANTUML,
            content=chart_element.content,
            language="plantuml",
            config=chart_element.attributes
        )
        
        # 检查是否能渲染
        if not self.plantuml_engine.can_render(chart_info):
            self.logger.warning("PlantUML引擎无法渲染此图表")
            self._add_chart_placeholder_fallback(chart_element)
            return
        
        # 生成唯一的文件名
        import hashlib
        content_hash = hashlib.md5(chart_element.content.encode()).hexdigest()[:8]
        chart_filename = f"plantuml_{content_hash}.png"
        chart_path = self.chart_cache_dir / chart_filename
        
        try:
            # 渲染图表
            success, rendered_path, error = self.plantuml_engine.render(chart_info, chart_path)
            
            if success and rendered_path and rendered_path.exists():
                # 插入图表到文档
                self._insert_chart_image(rendered_path, chart_element)
            else:
                self.logger.warning("PlantUML图表渲染结果无效")
                self._add_chart_placeholder_fallback(chart_element)
                
        except Exception as e:
            self.logger.error(f"PlantUML图表渲染失败: {e}")
            self._add_chart_placeholder_fallback(chart_element)
    
    def _insert_chart_image(self, image_path: Path, chart_element: ChartElement):
        """插入图表图片到文档
        
        Args:
            image_path: 图片路径
            chart_element: 图表元素
        """
        # 添加图表标题（如果有）
        if hasattr(chart_element, 'title') and chart_element.title:
            title_paragraph = self.document.add_paragraph(chart_element.title)
            title_paragraph.runs[0].font.bold = True
            title_paragraph.alignment = 1  # 居中对齐
        elif chart_element.chart_type:
            # 添加默认标题
            title_paragraph = self.document.add_paragraph(f"{chart_element.chart_type.title()} 图表")
            title_paragraph.runs[0].font.bold = True
            title_paragraph.alignment = 1  # 居中对齐
        
        # 插入图片
        try:
            # 获取图片尺寸配置
            max_width = self.config.get('chart.max_width', 6.0)  # 英寸
            
            # 添加图片
            paragraph = self.document.add_paragraph()
            paragraph.alignment = 1  # 居中对齐
            
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(max_width))
            
            self.logger.debug(f"成功插入图表图片: {image_path}")
            
        except Exception as e:
            self.logger.error(f"插入图片失败: {e}")
            # 降级为占位符
            self._add_chart_placeholder_fallback(chart_element)
    
    def _add_chart_placeholder_fallback(self, chart_element: ChartElement):
        """添加图表占位符（降级方案）
        
        Args:
            chart_element: 图表元素
        """
        self.logger.debug(f"使用占位符显示图表: {chart_element.chart_type}")
        
        # 添加图表标识
        chart_paragraph = self.document.add_paragraph(f"[{chart_element.chart_type.upper()} 图表]")
        chart_paragraph.runs[0].font.bold = True
        chart_paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色
        
        # 添加代码内容（用于调试）
        if self.config.get('debug.show_chart_code', False):
            code_paragraph = self.document.add_paragraph(chart_element.content)
            if 'Code' in self.document.styles:
                code_paragraph.style = 'Code'
        
        # 添加占位符说明
        placeholder_paragraph = self.document.add_paragraph("(图表渲染不可用，显示为占位符)")
        placeholder_paragraph.runs[0].font.italic = True
    
    def save_document(self, document, output_path: Path):
        """保存文档
        
        Args:
            document: Word文档对象
            output_path: 输出路径
        """
        try:
            # 确保输出目录存在
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 保存文档
            document.save(str(output_path))
            self.logger.info(f"文档保存成功: {output_path}")
            
        except Exception as e:
            self.logger.error(f"保存文档失败: {e}")
            raise
    
    def apply_template(self, template_name: str = "standard"):
        """应用文档模板
        
        Args:
            template_name: 模板名称
        """
        # TODO: 在后续任务中实现模板系统
        self.logger.info(f"应用模板: {template_name} (待实现)")
    
    def get_document_stats(self) -> Dict[str, Any]:
        """获取文档统计信息
        
        Returns:
            文档统计信息
        """
        if not self.document:
            return {}
        
        paragraphs = len(self.document.paragraphs)
        tables = len(self.document.tables)
        
        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "sections": len(self.document.sections)
        }
    
    def cleanup_chart_cache(self):
        """清理图表缓存
        """
        try:
            if self.chart_cache_dir.exists():
                for file in self.chart_cache_dir.glob("*"):
                    if file.is_file():
                        file.unlink()
                self.logger.debug("图表缓存已清理")
        except Exception as e:
            self.logger.warning(f"清理图表缓存失败: {e}")
    
    def __del__(self):
        """析构函数，清理资源"""
        try:
            # 可选择是否在对象销毁时清理缓存
            if self.config.get('chart.auto_cleanup', True):
                self.cleanup_chart_cache()
        except Exception:
            pass  # 忽略析构函数中的错误
