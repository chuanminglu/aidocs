"""
创建增强Word功能测试文档
"""
import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

def create_sample_images():
    """创建示例图片（如果不存在）"""
    from PIL import Image, ImageDraw
    
    temp_dir = Path("temp_images")
    temp_dir.mkdir(exist_ok=True)
    
    # 创建第一张图片 - 蓝色矩形
    img1 = Image.new('RGB', (300, 200), 'white')
    draw1 = ImageDraw.Draw(img1)
    draw1.rectangle([20, 20, 280, 180], fill='blue', outline='black')
    draw1.text((50, 90), "Sample Image 1", fill='white')
    img1_path = temp_dir / "sample1.png"
    img1.save(img1_path)
    
    # 创建第二张图片 - 红色圆形
    img2 = Image.new('RGB', (250, 250), 'white')
    draw2 = ImageDraw.Draw(img2)
    draw2.ellipse([25, 25, 225, 225], fill='red', outline='black')
    draw2.text((85, 115), "Sample Image 2", fill='white')
    img2_path = temp_dir / "sample2.png"
    img2.save(img2_path)
    
    return str(img1_path), str(img2_path)

def create_enhanced_test_document():
    """创建包含各种功能的测试文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('AI文档管理系统 - 增强功能测试', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加介绍段落
    intro = doc.add_paragraph()
    intro.add_run('本文档用于测试AI文档管理系统的')
    intro.add_run('增强Word解析功能').bold = True
    intro.add_run('，包括以下主要特性：')
    
    # 添加功能列表
    doc.add_paragraph('图片处理和提取功能', style='List Bullet')
    doc.add_paragraph('复杂表格支持（合并单元格等）', style='List Bullet')  
    doc.add_paragraph('样式保持功能', style='List Bullet')
    doc.add_paragraph('大纲结构解析', style='List Bullet')
    
    # 添加图片测试部分
    doc.add_heading('图片功能测试', level=1)
    
    try:
        # 创建示例图片
        img1_path, img2_path = create_sample_images()
        
        # 添加第一张图片
        doc.add_paragraph('这是第一张测试图片（蓝色矩形）：')
        doc.add_picture(img1_path, width=Inches(3.0))
        
        # 添加第二张图片
        doc.add_paragraph('这是第二张测试图片（红色圆形）：')
        doc.add_picture(img2_path, width=Inches(2.5))
        
        # 清理临时图片
        os.remove(img1_path)
        os.remove(img2_path)
        os.rmdir(Path(img1_path).parent)
        
        print("✅ 成功添加了2张测试图片")
        
    except Exception as e:
        print(f"⚠️  添加图片失败: {e}")
        doc.add_paragraph('图片添加失败，请检查PIL依赖是否安装正确')
    
    # 添加表格部分
    doc.add_heading('表格功能测试', level=1)
    
    # 创建复杂表格
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['项目名称', '开始日期', '结束日期', '负责人', '状态']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 数据行
    data = [
        ['AI文档管理', '2025-01-01', '2025-06-30', '张三', '进行中'],
        ['Word解析增强', '2025-02-01', '2025-03-15', '李四', '已完成'],
        ['UI界面优化', '2025-03-01', '2025-04-30', '王五', '计划中']
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.cell(i, j).text = cell_data
    
    # 添加样式测试部分
    doc.add_heading('样式功能测试', level=1)
    
    # 不同样式的文本
    p1 = doc.add_paragraph()
    p1.add_run('粗体文本').bold = True
    p1.add_run(' | ')
    p1.add_run('斜体文本').italic = True
    p1.add_run(' | ')
    p1.add_run('下划线文本').underline = True
    
    # 引用样式
    quote = doc.add_paragraph('这是一个引用样式的段落，展示不同的段落格式和样式保持功能。')
    quote.style = 'Quote'
    
    # 添加二级标题
    doc.add_heading('列表和编号测试', level=2)
    
    # 编号列表
    doc.add_paragraph('第一项：系统架构设计', style='List Number')
    doc.add_paragraph('第二项：功能模块实现', style='List Number')
    doc.add_paragraph('第三项：测试和部署', style='List Number')
    
    # 多级列表
    doc.add_paragraph('主要功能模块：', style='List Bullet')
    doc.add_paragraph('Word文档解析器', style='List Bullet 2')
    doc.add_paragraph('图片处理模块', style='List Bullet 2')
    doc.add_paragraph('表格解析模块', style='List Bullet 2')
    
    # 保存文档
    filename = 'enhanced_test_document.docx'
    doc.save(filename)
    print(f'✅ 创建增强测试文档: {filename}')
    return filename

if __name__ == "__main__":
    create_enhanced_test_document()
